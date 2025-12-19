import os
import re
import uuid
import tempfile
import subprocess
import base64
import logging
import gc
import time
import glob
import threading
import json
import traceback
from datetime import datetime, date, timedelta
from collections import defaultdict
from flask import Flask, render_template, request, jsonify, send_file, after_this_request, redirect, url_for, flash
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
import yt_dlp
import requests
import speech_recognition as sr
from pydub import AudioSegment
import pytesseract
from PIL import Image
from fuzzywuzzy import fuzz
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import db, AdminUser, AIProviderState, AIUsageLog, DailyStats, ActiveSession, ActivityLog, ToolStats, HourlyStats, ErrorLog
from services.ai_providers import ai_manager, AIProviderError, RateLimitError, ProviderNotConfiguredError

# --- كود إنشاء ملف الكوكيز تلقائياً من إعدادات السيرفر ---
# هذا يحمي حسابك بدلاً من رفع الملف على GitHub
cookie_content = os.environ.get('COOKIE_CONTENT')
if cookie_content:
    with open('cookies.txt', 'w') as f:
        f.write(cookie_content)
    print("✅ Cookies file created successfully from environment variables.")
else:
    print("⚠️ Warning: COOKIE_CONTENT not found in environment variables.")
# -------------------------------------------------------

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =============================================================================
# MEMORY OPTIMIZATION: Centralized cleanup utilities
# =============================================================================
TEMP_FILE_PATTERNS = [
    'audio_*.mp3', 'audio_*.wav', 'audio_*.m4a', 'audio_*.webm',
    'video_*.mp4', 'video_*.webm', 'video_*.mkv',
    'ocr_*.*', 'cut_*.*', 'arkan_*.*',
    '*_compressed.mp3', '*_chunk_*.wav'
]
TEMP_FILE_MAX_AGE_SECONDS = 3600
last_cleanup_time = 0

def cleanup_temp_files(force=False):
    """
    Clean up old temporary files to prevent disk space issues.
    
    Args:
        force: If True, clean all matching files regardless of age.
               If False, only clean files older than TEMP_FILE_MAX_AGE_SECONDS.
    """
    global last_cleanup_time
    current_time = time.time()
    
    if not force and current_time - last_cleanup_time < 300:
        return
    
    last_cleanup_time = current_time
    temp_dir = tempfile.gettempdir()
    cleaned_count = 0
    cleaned_size = 0
    
    for pattern in TEMP_FILE_PATTERNS:
        for filepath in glob.glob(os.path.join(temp_dir, pattern)):
            try:
                file_age = current_time - os.path.getmtime(filepath)
                if force or file_age > TEMP_FILE_MAX_AGE_SECONDS:
                    file_size = os.path.getsize(filepath)
                    os.remove(filepath)
                    cleaned_count += 1
                    cleaned_size += file_size
            except Exception as e:
                logger.debug(f"Could not clean {filepath}: {e}")
    
    if cleaned_count > 0:
        logger.info(f"Cleaned {cleaned_count} temp files, freed {cleaned_size / (1024*1024):.2f} MB")

def run_garbage_collection():
    """Run garbage collection to free memory after heavy operations."""
    gc.collect()
    logger.debug("Garbage collection completed")

def cleanup_after_request():
    """Combined cleanup: temp files + garbage collection."""
    cleanup_temp_files()
    run_garbage_collection()

def safe_remove_file(filepath):
    """Safely remove a file with error handling."""
    if filepath and os.path.exists(filepath):
        try:
            os.remove(filepath)
            return True
        except Exception as e:
            logger.debug(f"Could not remove {filepath}: {e}")
    return False

def safe_remove_files(*filepaths):
    """Safely remove multiple files."""
    for filepath in filepaths:
        safe_remove_file(filepath)

# =============================================================================
# HELPER FUNCTION: Centralized YouTube Audio Download
# This function handles all YouTube downloads with proper error handling,
# cookies support, and bypass techniques for 403 errors.
# =============================================================================
def download_audio_from_youtube(url: str, output_dir: str = None) -> str:
    """
    Download audio from YouTube URL using yt-dlp with robust error handling.
    
    Args:
        url: YouTube video URL
        output_dir: Directory to save the file (uses temp dir if None)
    
    Returns:
        Path to the downloaded audio file
    
    Raises:
        Exception: If download fails
    """
    if output_dir is None:
        output_dir = tempfile.gettempdir()
    
    out_template = os.path.join(output_dir, "%(id)s.%(ext)s")
    
    # Try multiple format strategies to handle various YouTube restrictions
    format_strategies = [
        "bestaudio[ext=m4a]/bestaudio[ext=webm]/bestaudio",
        "bestaudio/best",
        "best[ext=m4a]/best[ext=webm]/best",
        "worst[ext=m4a]/worst",
    ]
    
    for format_str in format_strategies:
        try:
            logger.info(f"Attempting YouTube download with format: {format_str}")
            ydl_opts = {
                "format": format_str,
                "outtmpl": out_template,
                "noplaylist": True,
                "quiet": False,
                "no_warnings": False,
                "nocheckcertificate": True,
                "geo_bypass": True,
                "geo_bypass_country": "US",
                "socket_timeout": 300,
                "retries": 3,
                "fragment_retries": 3,
                "concurrent_fragment_downloads": 4,
                "buffersize": 1024 * 16,
                "http_chunk_size": 10485760,
                "allow_unplayable_formats": True,
                "skip_unavailable_fragments": True,
                "http_headers": {
                    "User-Agent": (
                        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) "
                        "Chrome/120.0.0.0 Safari/537.36"
                    ),
                    "Accept-Language": "en-US,en;q=0.5",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Referer": "https://www.youtube.com/",
                },
                "postprocessors": [{
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "mp3",
                    "preferredquality": "192",
                }],
            }
            
            # Add cookies file if available (critical for bypassing YouTube restrictions)
            if os.path.exists('cookies.txt'):
                ydl_opts['cookiefile'] = 'cookies.txt'
                logger.info("Using cookies.txt for YouTube authentication")
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)
                filename = ydl.prepare_filename(info)
                # Convert to mp3 if needed
                base_name = os.path.splitext(filename)[0]
                mp3_file = base_name + '.mp3'
                
                if os.path.exists(mp3_file):
                    logger.info(f"Downloaded YouTube audio to: {mp3_file}")
                    return mp3_file
                elif os.path.exists(filename):
                    logger.info(f"Downloaded YouTube audio to: {filename}")
                    return filename
        except Exception as e:
            logger.warning(f"Format strategy '{format_str}' failed: {e}")
            continue
    
    # If all strategies fail, raise an error
    raise Exception("فشل في تحميل الصوت من YouTube بعد عدة محاولات. قد يكون الفيديو محميًا أو غير متاح في منطقتك.")

app = Flask(__name__, static_folder='static', static_url_path='/static')
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1)
app.secret_key = os.environ.get("SESSION_SECRET") or "dev-secret-key-change-in-production"
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024 * 1024

database_url = os.environ.get("DATABASE_URL") or os.environ.get("SUPABASE_URL")
if not database_url:
    raise RuntimeError(
        "DATABASE_URL or SUPABASE_URL environment variable is not set. "
        "Please set it to your PostgreSQL connection string. "
        "Example: postgresql://user:password@host:port/database"
    )
app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'admin_login'


@login_manager.user_loader
def load_user(user_id):
    return AdminUser.query.get(int(user_id))


@app.after_request
def after_request_cleanup(response):
    """Periodic cleanup after requests to manage memory and disk space."""
    cleanup_after_request()
    return response


def parse_user_agent(user_agent_string):
    """Parse user agent string to extract device type, browser, and OS."""
    if not user_agent_string:
        return 'unknown', 'unknown', 'unknown'
    
    ua = user_agent_string.lower()
    
    if 'mobile' in ua or 'android' in ua or 'iphone' in ua:
        device_type = 'mobile'
    elif 'tablet' in ua or 'ipad' in ua:
        device_type = 'tablet'
    else:
        device_type = 'desktop'
    
    browser = 'unknown'
    if 'chrome' in ua and 'edg' not in ua:
        browser = 'Chrome'
    elif 'firefox' in ua:
        browser = 'Firefox'
    elif 'safari' in ua and 'chrome' not in ua:
        browser = 'Safari'
    elif 'edg' in ua:
        browser = 'Edge'
    elif 'opera' in ua:
        browser = 'Opera'
    
    os_name = 'unknown'
    if 'windows' in ua:
        os_name = 'Windows'
    elif 'mac os' in ua or 'macos' in ua:
        os_name = 'macOS'
    elif 'linux' in ua:
        os_name = 'Linux'
    elif 'android' in ua:
        os_name = 'Android'
    elif 'iphone' in ua or 'ipad' in ua:
        os_name = 'iOS'
    
    return device_type, browser, os_name


SESSION_INACTIVE_MINUTES = 5
SESSION_CLEANUP_INTERVAL = 300
last_session_cleanup = 0


@app.before_request
def track_active_session():
    """Track active sessions for analytics."""
    global last_session_cleanup
    
    if request.endpoint in ('static', 'admin_api_stats', 'admin_api_active_sessions'):
        return
    
    try:
        session_id = request.cookies.get('arkan_session')
        if not session_id:
            session_id = str(uuid.uuid4())
            request.new_session_id = session_id
        
        ip_address = request.remote_addr or 'unknown'
        user_agent = request.headers.get('User-Agent', '')
        device_type, browser, os_name = parse_user_agent(user_agent)
        
        existing_session = ActiveSession.query.filter_by(session_id=session_id).first()
        
        if existing_session:
            existing_session.last_seen = datetime.utcnow()
            existing_session.page_views += 1
            existing_session.is_active = True
        else:
            new_session = ActiveSession(
                session_id=session_id,
                ip_address=ip_address,
                user_agent=user_agent[:500] if user_agent else None,
                device_type=device_type,
                browser=browser,
                os_name=os_name
            )
            db.session.add(new_session)
        
        current_time = time.time()
        if current_time - last_session_cleanup > SESSION_CLEANUP_INTERVAL:
            inactive_threshold = datetime.utcnow() - timedelta(minutes=SESSION_INACTIVE_MINUTES)
            ActiveSession.query.filter(ActiveSession.last_seen < inactive_threshold).update({'is_active': False})
            old_threshold = datetime.utcnow() - timedelta(days=7)
            ActiveSession.query.filter(ActiveSession.last_seen < old_threshold).delete()
            last_session_cleanup = current_time
        
        db.session.commit()
    except Exception as e:
        logger.debug(f"Session tracking error: {e}")
        try:
            db.session.rollback()
        except:
            pass


@app.after_request
def set_session_cookie(response):
    """Set session cookie for new visitors."""
    new_session_id = getattr(request, 'new_session_id', None)
    if new_session_id:
        response.set_cookie(
            'arkan_session',
            new_session_id,
            max_age=60*60*24*365,
            httponly=True,
            samesite='Lax'
        )
    return response


UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
ALLOWED_AUDIO_EXTENSIONS = {'mp3', 'wav', 'ogg', 'm4a', 'flac', 'aac', 'wma'}

ANIME_SIMILARITY_THRESHOLD = 0.85
PODCAST_NAME_SIMILARITY_THRESHOLD = 90

# =============================================================================
# ACTIVITY LOGGING HELPER
# =============================================================================
# Logs all tool activities (AI and non-AI) to the ActivityLog table
# for the unified admin dashboard.
# =============================================================================

def log_activity(tool_name, action, status, duration_ms=None, file_size=None, 
                 error_message=None, details=None):
    """
    Log tool activity to the database for admin dashboard.
    
    Args:
        tool_name: Name of the tool (e.g., 'video_transcription', 'pdf_converter')
        action: Action performed (e.g., 'transcribe', 'convert', 'download')
        status: Status of the operation ('success', 'error', 'warning')
        duration_ms: How long the operation took in milliseconds
        file_size: Size of the processed file in bytes
        error_message: Error message if status is 'error'
        details: Additional details as JSON string
    """
    try:
        session_id = request.cookies.get('arkan_session', 'unknown')
        ip_address = request.remote_addr or 'unknown'
        user_agent = request.headers.get('User-Agent', '')
        device_type, _, _ = parse_user_agent(user_agent)
        
        activity = ActivityLog(
            session_id=session_id,
            ip_address=ip_address,
            tool_name=tool_name,
            action=action,
            status=status,
            duration_ms=duration_ms,
            file_size=file_size,
            error_message=error_message[:500] if error_message else None,
            details=details,
            user_agent=user_agent[:500] if user_agent else None,
            device_type=device_type
        )
        db.session.add(activity)
        
        # Update ToolStats
        today = date.today()
        tool_stat = ToolStats.query.filter_by(date=today, tool_name=tool_name).first()
        if tool_stat:
            tool_stat.usage_count += 1
            if status == 'success':
                tool_stat.success_count += 1
            else:
                tool_stat.error_count += 1
            if duration_ms:
                # Calculate new average
                total_duration = tool_stat.avg_duration_ms * (tool_stat.usage_count - 1) + duration_ms
                tool_stat.avg_duration_ms = total_duration / tool_stat.usage_count
            if file_size:
                tool_stat.total_file_size += file_size
        else:
            tool_stat = ToolStats(
                date=today,
                tool_name=tool_name,
                usage_count=1,
                success_count=1 if status == 'success' else 0,
                error_count=1 if status != 'success' else 0,
                avg_duration_ms=duration_ms or 0,
                total_file_size=file_size or 0
            )
            db.session.add(tool_stat)
        
        db.session.commit()
    except Exception as e:
        logger.debug(f"Activity logging error: {e}")
        try:
            db.session.rollback()
        except:
            pass


def log_error(error_type, error_message, stack_trace=None, provider=None, 
              tool_name=None, request_data=None):
    """
    Log errors to the ErrorLog table for debugging.
    
    Args:
        error_type: Type of error (e.g., 'AIError', 'FileError', 'ValidationError')
        error_message: The error message
        stack_trace: Full stack trace if available
        provider: AI provider name if applicable
        tool_name: Name of the tool where error occurred
        request_data: Request data for debugging
    """
    try:
        session_id = request.cookies.get('arkan_session', 'unknown')
        
        error_log = ErrorLog(
            error_type=error_type,
            error_message=error_message[:1000] if error_message else 'Unknown error',
            stack_trace=stack_trace,
            provider=provider,
            tool_name=tool_name,
            session_id=session_id,
            request_data=request_data
        )
        db.session.add(error_log)
        db.session.commit()
    except Exception as e:
        logger.debug(f"Error logging failed: {e}")
        try:
            db.session.rollback()
        except:
            pass


# =============================================================================
# AI MANAGER WRAPPER FUNCTIONS
# =============================================================================
# Provides backward-compatible wrapper functions using the new AIManager
# from services/ai_providers.py which uses Groq + HuggingFace instead of Gemini.
# =============================================================================

def is_ai_provider_configured():
    """Check if any AI provider is configured (backward compatibility for GEMINI_KEYS checks)."""
    return ai_manager.groq.is_configured or ai_manager.huggingface.is_configured

def get_available_keys_count():
    """Get count of available AI providers (backward compatibility)."""
    count = 0
    if ai_manager.groq.is_configured:
        count += 1
    if ai_manager.huggingface.is_configured:
        count += 1
    return count

GEMINI_KEYS = is_ai_provider_configured()

MAX_AUDIO_DURATION_MINUTES = 120
MAX_LONG_VIDEO_PER_SESSION = 3

def check_audio_duration_limit(audio_path):
    """
    Check if audio/video duration is within allowed limits (2 hours max).
    
    Returns:
        tuple: (allowed: bool, duration_minutes: float, error_message: str or None)
    """
    try:
        audio = AudioSegment.from_file(audio_path)
        duration_minutes = len(audio) / 60000
        
        if duration_minutes > MAX_AUDIO_DURATION_MINUTES:
            hours = duration_minutes / 60
            max_hours = MAX_AUDIO_DURATION_MINUTES / 60
            return False, duration_minutes, (
                f"مدة الملف ({hours:.1f} ساعة / {duration_minutes:.0f} دقيقة) أكبر من الحد الأقصى المسموح "
                f"({max_hours:.0f} ساعة). يُنصح بتقسيم الملف إلى أجزاء أقصر."
            )
        
        return True, duration_minutes, None
        
    except Exception as e:
        logging.error(f"[AIManager] Could not check audio duration: {e}")
        return True, 0, None


def call_gemini_text(prompt, max_retries=None, operation='text'):
    """
    Call LLM for text-only prompts - backward compatible wrapper.
    Now uses Groq Llama instead of Gemini.
    
    Args:
        prompt: The text prompt to send to the model
        max_retries: Ignored (kept for backward compatibility)
        operation: Operation name for logging
    
    Returns:
        str: Response text or None if failed
    """
    try:
        return ai_manager.call_llm(prompt)
    except AIProviderError as e:
        logging.warning(f"[call_gemini_text] {e}")
        return None


def call_gemini_vision(image_or_path, prompt, max_retries=None, operation='vision'):
    """
    Call Vision API - backward compatible wrapper.
    Now uses HuggingFace Vision instead of Gemini.
    
    Args:
        image_or_path: Path to image file (PIL Image not supported with HuggingFace)
        prompt: The prompt/question to send to the model
        max_retries: Ignored (kept for backward compatibility)
        operation: Operation name for logging
    
    Returns:
        str: Response text or None if failed
    """
    try:
        # If it's a PIL Image, save it temporarily
        if hasattr(image_or_path, 'save'):
            import tempfile
            temp_path = tempfile.mktemp(suffix='.png')
            image_or_path.save(temp_path)
            try:
                return ai_manager.analyze_image(temp_path, question=prompt)
            finally:
                safe_remove_file(temp_path)
        else:
            return ai_manager.analyze_image(image_or_path, question=prompt)
    except AIProviderError as e:
        logging.warning(f"[call_gemini_vision] {e}")
        return None


def allowed_image(filename):
    return '.' in filename and filename.rsplit(
        '.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


def identify_podcast_with_gemini(image_path):
    """Use Gemini AI Vision to identify podcast from image"""
    prompt = """Analyze this image carefully. It could be a YouTube video screenshot, podcast artwork, or a photo of podcast hosts.
    
    Identify the Podcast Name, Channel Name, Show Name, or Host Names based on:
    - Visible logos, titles, or text
    - Recognizable faces of podcast hosts or celebrities
    - Studio setup, microphones, or recording environment
    - Channel branding or show graphics
    - Any visual elements that identify the content
    
    Return the response in this format: PODCAST_NAME|HOST_NAMES|PLATFORM
    - PODCAST_NAME: The main podcast/show/channel name
    - HOST_NAMES: Names of visible hosts (or 'Unknown' if not recognized)
    - PLATFORM: Where this podcast is likely found (YouTube, Spotify, Apple Podcasts, or General)
    
    If you cannot identify it at all, return 'UNKNOWN'.
    Be specific and accurate. Arabic and English names are both acceptable.
    
    Example: The Joe Rogan Experience|Joe Rogan|Spotify
    Example: بودكاست فنجان|عبدالرحمن أبومالح|YouTube"""

    result = call_gemini_vision(image_path, prompt)

    if result and result.upper() != 'UNKNOWN' and len(result) > 2:
        return result
    return None


def identify_podcast_from_transcript(transcript):
    """Use Gemini AI to identify podcast from transcribed audio text"""
    prompt = f"""Based on this audio transcript from a podcast or show, identify the Podcast name, Host, or Show name.
    
Transcript:
{transcript[:2000]}

Instructions:
- Look for any mentions of podcast names, show names, host introductions, channel names
- Consider common podcast phrases like "welcome to...", "this is...", "you're listening to..."
- Return ONLY the podcast/show name, nothing else
- If you cannot identify it, return 'UNKNOWN'
- Be specific and accurate."""

    result = call_gemini_text(prompt)

    if result and result.upper() != 'UNKNOWN' and len(result) > 2:
        clean_result = result.replace('"', '').replace("'", '').strip()
        if clean_result.upper() != 'UNKNOWN':
            return clean_result
    return None


def compress_audio_for_upload(audio_path):
    """Compress audio to low-bitrate MP3 for faster upload to Gemini"""
    try:
        audio = AudioSegment.from_file(audio_path)
        audio = audio.set_channels(1)
        audio = audio.set_frame_rate(16000)

        compressed_path = audio_path.rsplit('.', 1)[0] + '_compressed.mp3'
        audio.export(compressed_path, format="mp3", bitrate="32k")

        logging.info(
            f"Audio compressed for upload: {audio_path} -> {compressed_path}")
        return compressed_path
    except Exception as e:
        logging.error(f"Audio compression failed: {e}")
        return audio_path


LANGUAGE_NAMES = {
    'ar': 'Arabic',
    'en': 'English',
    'fr': 'French',
    'es': 'Spanish',
    'de': 'German',
    'tr': 'Turkish',
    'ur': 'Urdu',
    'hi': 'Hindi',
    'id': 'Indonesian',
    'pt': 'Portuguese',
    'ru': 'Russian',
    'ja': 'Japanese',
    'ko': 'Korean',
    'zh': 'Chinese',
    'auto': 'Auto-detect'
}


def transcribe_audio_with_gemini(audio_path, language='ar'):
    """
    Transcribe audio file using Groq Whisper.
    
    Features:
    - Uses Groq Whisper for fast, accurate transcription
    - Duration limit: 2 hours max
    - Session-based rate limiting
    - Comprehensive logging for monitoring
    - Graceful error handling with Arabic messages
    """
    allowed, duration_minutes, error_msg = check_audio_duration_limit(audio_path)
    if not allowed:
        raise ValueError(error_msg)
    
    logging.info(f"[Transcription] Audio duration: {duration_minutes:.1f} minutes")
    
    compressed_path = compress_audio_for_upload(audio_path)
    upload_path = compressed_path if compressed_path != audio_path else audio_path
    
    try:
        lang_code = language if language != 'auto' else None
        result = ai_manager.transcribe_audio(upload_path, language=lang_code)
        
        if compressed_path != audio_path:
            safe_remove_file(compressed_path)
        
        return result
        
    except AIProviderError as e:
        if compressed_path != audio_path:
            safe_remove_file(compressed_path)
        logging.error(f"[Transcription] AI error: {e}")
        raise ValueError(str(e))
    except Exception as e:
        if compressed_path != audio_path:
            safe_remove_file(compressed_path)
        logging.error(f"[Transcription] Unexpected error: {e}")
        raise ValueError(f"فشل في تحويل الصوت إلى نص. الخطأ: {str(e)[:100]}")


def generate_podcast_search_links(podcast_name):
    """Generate search links for YouTube, Spotify, and SoundCloud"""
    if not podcast_name or not podcast_name.strip():
        return {'youtube': '', 'spotify': '', 'soundcloud': ''}
    from urllib.parse import quote
    clean_name = podcast_name.strip()
    encoded_name = quote(clean_name)
    return {
        'youtube':
        f"https://www.youtube.com/results?search_query={encoded_name}",
        'spotify': f"https://open.spotify.com/search/{encoded_name}",
        'soundcloud': f"https://soundcloud.com/search?q={encoded_name}"
    }


def identify_anime_with_gemini(image_path):
    """Use Gemini AI Vision to identify anime from image"""
    prompt = """Analyze this anime screenshot or image carefully.
    Identify the anime title based on the characters, art style, scene, character designs, backgrounds, or any recognizable elements.
    
    Instructions:
    - If you recognize the anime, return the title in this format: ANIME_NAME|EPISODE_INFO|DESCRIPTION
    - ANIME_NAME: The official English or Romaji title of the anime
    - EPISODE_INFO: Episode number if recognizable, or 'Unknown' if not
    - DESCRIPTION: Brief description of the scene or characters shown (1-2 sentences)
    - If you cannot identify it at all, return 'UNKNOWN'
    - Be as specific and accurate as possible
    - Consider popular anime series, movies, and OVAs
    
    Example response: Demon Slayer: Kimetsu no Yaiba|Episode 19|Tanjiro performing the Hinokami Kagura dance move against Rui"""

    result = call_gemini_vision(image_path, prompt)

    if result and result.upper() != 'UNKNOWN' and len(result) > 2:
        return result.replace('"', '').strip()
    return None


def identify_anime_by_description(description):
    """Use Gemini AI to identify anime from text description"""
    prompt = f"""أنت خبير في الأنمي. بناءً على هذا الوصف، حدد اسم الأنمي.

الوصف:
{description}

التعليمات:
- حلل الوصف للبحث عن أسماء الشخصيات، عناصر القصة، الإعدادات، القوى الخارقة، أو أي سمات مميزة
- ابحث عن أي تفاصيل يمكن أن تشير إلى أنمي معين مثل:
  * أسماء الشخصيات (بالعربية أو الإنجليزية أو اليابانية)
  * وصف المظهر أو الملابس المميزة
  * القوى الخاصة أو الأسلحة
  * الأحداث أو المشاهد المميزة
  * أسلوب الرسم أو الأنيميشن
  * أسماء الأماكن أو العوالم
- أعد الرد بهذا الشكل: ANIME_NAME|CONFIDENCE|ALTERNATIVES
- ANIME_NAME: الاسم الأكثر احتمالاً للأنمي (بالإنجليزية أو الرومانجي)
- CONFIDENCE: High أو Medium أو Low
- ALTERNATIVES: أنميات بديلة محتملة مفصولة بفواصل (حتى 5)
- إذا لم تتمكن من التعرف عليه، أعد 'UNKNOWN'

أمثلة:
- "شخص يقاتل عمالقة وجدران ضخمة" -> Attack on Titan|High|Kabaneri of the Iron Fortress, God Eater
- "ولد شعره أشقر يريد أن يصبح هوكاجي" -> Naruto|High|Boruto
- "قراصنة يبحثون عن كنز" -> One Piece|High|Black Lagoon, Pirates of the Caribbean
- "شخص يستخدم دفتر لقتل الناس" -> Death Note|High|Future Diary"""

    result = call_gemini_text(prompt)

    if result and result.upper() != 'UNKNOWN' and len(result) > 2:
        return result.strip()
    return None


def allowed_audio(filename):
    return '.' in filename and filename.rsplit(
        '.', 1)[1].lower() in ALLOWED_AUDIO_EXTENSIONS


def validate_time_format(time_str):
    if not time_str:
        return None

    time_str = time_str.strip()
    parts = time_str.split(':')

    try:
        if len(parts) == 2:
            minutes, seconds = int(parts[0]), int(parts[1])
            if minutes < 0 or seconds < 0 or seconds >= 60:
                return None
            return minutes * 60 + seconds
        elif len(parts) == 3:
            hours, minutes, seconds = int(parts[0]), int(parts[1]), int(
                parts[2])
            if hours < 0 or minutes < 0 or minutes >= 60 or seconds < 0 or seconds >= 60:
                return None
            return hours * 3600 + minutes * 60 + seconds
        else:
            return None
    except ValueError:
        return None


@app.route('/')
def index():
    return render_template('index.html',
                         supabase_url=os.environ.get('SUPABASE_URL', ''),
                         supabase_anon_key=os.environ.get('SUPABASE_ANON_KEY', ''))


@app.route('/logo.png')
def serve_logo():
    return send_file('logo.png', mimetype='image/png')


@app.route('/video-info', methods=['POST'])
def video_info():
    try:
        data = request.get_json()
        url = data.get('url', '')

        if not url:
            return jsonify({'error': 'الرجاء إدخال رابط الفيديو'}), 400

        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'skip_download': True,
            'noplaylist': True,
            'force_generic_extractor': False,
            'cookiefile':
            'cookies.txt' if os.path.exists('cookies.txt') else None,
            'socket_timeout': 1800,
            'retries': 3,
            'http_headers': {
                'User-Agent':
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept':
                'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-us,en;q=0.5',
                'Sec-Fetch-Mode': 'navigate',
            },
        }

        if ydl_opts.get('cookiefile') is None:
            del ydl_opts['cookiefile']

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)

            if not info:
                return jsonify({'error':
                                'لم يتم العثور على معلومات الفيديو'}), 400

            duration_seconds = info.get('duration') or 0
            if duration_seconds:
                hours = duration_seconds // 3600
                minutes = (duration_seconds % 3600) // 60
                seconds = duration_seconds % 60
                if hours > 0:
                    duration_formatted = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                else:
                    duration_formatted = f"{minutes:02d}:{seconds:02d}"
            else:
                duration_formatted = "00:00"

            title = info.get('title') or info.get('fulltitle') or 'غير معروف'
            channel = info.get('uploader') or info.get('channel') or info.get(
                'uploader_id') or 'غير معروف'
            thumbnail = info.get('thumbnail') or ''
            if not thumbnail and info.get('thumbnails'):
                thumbnails = info.get('thumbnails', [])
                if thumbnails:
                    thumbnail = thumbnails[-1].get('url', '')

            return jsonify({
                'title': title,
                'duration': duration_formatted,
                'duration_seconds': duration_seconds,
                'thumbnail': thumbnail,
                'channel': channel
            })

    except Exception as e:
        logging.error(f"Video info error: {str(e)}")
        return jsonify({'error': f'خطأ في جلب معلومات الفيديو: {str(e)}'}), 400


# =============================================================================
# دالة جلب الجودات المتاحة من يوتيوب
# =============================================================================
@app.route('/get-video-formats', methods=['POST'])
def get_video_formats():
    """جلب قائمة الجودات والفورمات المتاحة للفيديو"""
    try:
        data = request.get_json()
        url = data.get('url', '')

        if not url:
            return jsonify({'error': 'الرجاء إدخال رابط الفيديو'}), 400

        logging.info(f"جلب الجودات المتاحة للفيديو: {url}")

        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'skip_download': True,
            'noplaylist': True,
            'socket_timeout': 1800,
            'retries': 3,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            },
        }

        if os.path.exists('cookies.txt'):
            ydl_opts['cookiefile'] = 'cookies.txt'

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)

            if not info:
                return jsonify({'error': 'لم يتم العثور على معلومات الفيديو'}), 400

            formats = info.get('formats', [])
            available_qualities = []
            seen_resolutions = set()

            # فلترة وتجميع الجودات المتاحة
            for fmt in formats:
                height = fmt.get('height')
                vcodec = fmt.get('vcodec', 'none')
                acodec = fmt.get('acodec', 'none')
                ext = fmt.get('ext', '')
                filesize = fmt.get('filesize') or fmt.get('filesize_approx') or 0

                # تجاهل الفورمات بدون ارتفاع محدد
                if height and vcodec != 'none':
                    resolution = f"{height}p"
                    if resolution not in seen_resolutions:
                        seen_resolutions.add(resolution)
                        size_mb = round(filesize / (1024 * 1024), 1) if filesize else 0
                        available_qualities.append({
                            'resolution': resolution,
                            'height': height,
                            'type': 'video',
                            'size_mb': size_mb,
                            'format': ext
                        })

            # ترتيب الجودات من الأعلى للأقل
            available_qualities.sort(key=lambda x: x['height'], reverse=True)

            # إضافة خيار الصوت فقط
            audio_formats = [f for f in formats if f.get('acodec') != 'none' and f.get('vcodec') == 'none']
            best_audio = None
            for af in audio_formats:
                if af.get('abr'):
                    if not best_audio or af.get('abr', 0) > best_audio.get('abr', 0):
                        best_audio = af

            # جلب معلومات الفيديو الأساسية
            duration_seconds = info.get('duration') or 0
            title = info.get('title') or 'غير معروف'
            thumbnail = info.get('thumbnail') or ''

            logging.info(f"تم جلب {len(available_qualities)} جودة متاحة")

            return jsonify({
                'success': True,
                'title': title,
                'duration_seconds': duration_seconds,
                'thumbnail': thumbnail,
                'qualities': available_qualities,
                'has_audio': best_audio is not None,
                'default_quality': available_qualities[0]['resolution'] if available_qualities else '720p'
            })

    except Exception as e:
        logging.error(f"خطأ في جلب الجودات: {str(e)}")
        return jsonify({'error': f'خطأ في جلب الجودات المتاحة: {str(e)}'}), 400


# =============================================================================
# دالة تحميل الفيديو/الصوت من يوتيوب
# =============================================================================
def download_youtube_media(url: str, quality: str, download_type: str, output_dir: str) -> str:
    """
    تحميل الفيديو أو الصوت من يوتيوب بالجودة المحددة
    
    Args:
        url: رابط الفيديو
        quality: الجودة المطلوبة (مثل 720p, 1080p)
        download_type: نوع التحميل (video أو audio)
        output_dir: مجلد الحفظ
    
    Returns:
        مسار الملف المحمل
    """
    unique_id = str(uuid.uuid4())[:8]
    
    if download_type == 'audio':
        # تحميل الصوت فقط
        output_template = os.path.join(output_dir, f'audio_{unique_id}.%(ext)s')
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': output_template,
            'noplaylist': True,
            'quiet': True,
            'no_warnings': True,
            'socket_timeout': 7200,
            'retries': 5,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            },
        }
    else:
        # تحميل الفيديو بالجودة المحددة
        height = quality.replace('p', '') if quality else '720'
        output_template = os.path.join(output_dir, f'video_{unique_id}.%(ext)s')
        
        # اختيار أفضل فورمات بناءً على الجودة المطلوبة
        format_string = f'bestvideo[height<={height}][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height<={height}]+bestaudio/best[height<={height}]/best'
        
        ydl_opts = {
            'format': format_string,
            'merge_output_format': 'mp4',
            'outtmpl': output_template,
            'noplaylist': True,
            'quiet': True,
            'no_warnings': True,
            'socket_timeout': 7200,
            'retries': 5,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            },
        }

    if os.path.exists('cookies.txt'):
        ydl_opts['cookiefile'] = 'cookies.txt'

    logging.info(f"بدء تحميل {download_type} بجودة {quality}")
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])

    # البحث عن الملف المحمل
    base_name = output_template.replace('.%(ext)s', '')
    for ext in ['mp4', 'mp3', 'mkv', 'webm', 'm4a', 'wav']:
        possible_file = f"{base_name}.{ext}"
        if os.path.exists(possible_file):
            logging.info(f"تم تحميل الملف: {possible_file}")
            return possible_file

    raise Exception("لم يتم العثور على الملف المحمل")


# =============================================================================
# دالة قص الملف باستخدام ffmpeg
# =============================================================================
def cut_media_segment(input_path: str, start_seconds: int, end_seconds: int, 
                      download_type: str, output_dir: str) -> str:
    """
    قص جزء من الفيديو أو الصوت باستخدام ffmpeg
    
    Args:
        input_path: مسار الملف الأصلي
        start_seconds: وقت البداية بالثواني
        end_seconds: وقت النهاية بالثواني
        download_type: نوع الملف (video أو audio)
        output_dir: مجلد الحفظ
    
    Returns:
        مسار الملف المقصوص
    """
    unique_id = str(uuid.uuid4())[:8]
    duration = end_seconds - start_seconds
    
    if download_type == 'audio':
        output_path = os.path.join(output_dir, f'clip_{unique_id}.mp3')
        # قص الصوت مع ترميز خفيف
        ffmpeg_cmd = [
            'ffmpeg', '-y',
            '-ss', str(start_seconds),
            '-i', input_path,
            '-t', str(duration),
            '-c:a', 'libmp3lame',
            '-b:a', '192k',
            '-ar', '44100',
            output_path
        ]
    else:
        output_path = os.path.join(output_dir, f'clip_{unique_id}.mp4')
        # قص الفيديو بدون إعادة ترميز قدر الإمكان (أسرع وأخف)
        ffmpeg_cmd = [
            'ffmpeg', '-y',
            '-ss', str(start_seconds),
            '-i', input_path,
            '-t', str(duration),
            '-c', 'copy',  # نسخ مباشر بدون إعادة ترميز
            '-avoid_negative_ts', 'make_zero',
            output_path
        ]

    logging.info(f"بدء قص الملف من {start_seconds}s إلى {end_seconds}s")
    
    try:
        result = subprocess.run(ffmpeg_cmd, check=True, capture_output=True, timeout=7200)
        
        # إذا فشل القص بدون ترميز، نحاول مع إعادة الترميز
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            if download_type == 'video':
                logging.info("القص بدون ترميز فشل، جاري المحاولة مع إعادة الترميز...")
                ffmpeg_cmd = [
                    'ffmpeg', '-y',
                    '-ss', str(start_seconds),
                    '-i', input_path,
                    '-t', str(duration),
                    '-c:v', 'libx264',
                    '-c:a', 'aac',
                    '-strict', 'experimental',
                    output_path
                ]
                subprocess.run(ffmpeg_cmd, check=True, capture_output=True, timeout=7200)
    except subprocess.CalledProcessError as e:
        # محاولة أخيرة مع إعادة الترميز الكامل
        if download_type == 'video':
            ffmpeg_cmd = [
                'ffmpeg', '-y',
                '-ss', str(start_seconds),
                '-i', input_path,
                '-t', str(duration),
                '-c:v', 'libx264',
                '-c:a', 'aac',
                '-strict', 'experimental',
                output_path
            ]
            subprocess.run(ffmpeg_cmd, check=True, capture_output=True, timeout=7200)
        else:
            raise e

    if os.path.exists(output_path):
        logging.info(f"تم قص الملف بنجاح: {output_path}")
        return output_path
    else:
        raise Exception("فشل في إنشاء الملف المقصوص")


@app.route('/process-video', methods=['POST'])
def process_video():
    """
    معالجة وقص الفيديو من يوتيوب مع دعم اختيار الجودة ونوع التحميل
    الحقول المدعومة:
    - url: رابط الفيديو (مطلوب)
    - start_time: وقت البداية (مطلوب)
    - end_time: وقت النهاية (مطلوب)
    - quality: الجودة (اختياري - افتراضي 720p)
    - download_type: نوع التحميل video أو audio (اختياري - افتراضي video)
    """
    start_time_log = time.time()
    temp_media = None
    output_file = None
    file_size = 0

    try:
        data = request.get_json()
        url = data.get('url', '')
        start_time = data.get('start_time', '00:00')
        end_time = data.get('end_time', '')
        quality = data.get('quality', '720p')  # الجودة الافتراضية
        download_type = data.get('download_type', 'video')  # النوع الافتراضي

        logging.info(f"طلب قص فيديو: quality={quality}, type={download_type}")

        if not url:
            return jsonify({'error': 'الرجاء إدخال رابط الفيديو'}), 400

        if not end_time:
            return jsonify({'error': 'الرجاء تحديد وقت النهاية'}), 400

        # التحقق من نوع التحميل
        if download_type not in ['video', 'audio']:
            download_type = 'video'

        start_seconds = validate_time_format(start_time)
        end_seconds = validate_time_format(end_time)

        if start_seconds is None:
            return jsonify({
                'error':
                'صيغة وقت البداية غير صحيحة. استخدم الصيغة MM:SS أو HH:MM:SS'
            }), 400

        if end_seconds is None:
            return jsonify({
                'error':
                'صيغة وقت النهاية غير صحيحة. استخدم الصيغة MM:SS أو HH:MM:SS'
            }), 400

        if end_seconds <= start_seconds:
            return jsonify(
                {'error': 'وقت النهاية يجب أن يكون أكبر من وقت البداية'}), 400

        # تحميل الملف باستخدام الدالة المساعدة الجديدة
        logging.info(f"بدء تحميل الملف من يوتيوب...")
        temp_media = download_youtube_media(url, quality, download_type, UPLOAD_FOLDER)
        logging.info(f"تم تحميل الملف: {temp_media}")

        # قص الملف باستخدام الدالة المساعدة
        logging.info(f"بدء قص الملف...")
        output_file = cut_media_segment(temp_media, start_seconds, end_seconds, 
                                         download_type, UPLOAD_FOLDER)
        logging.info(f"تم قص الملف: {output_file}")

        # حذف الملف الأصلي بعد القص
        safe_remove_file(temp_media)
        logging.info(f"تم حذف الملف المؤقت: {temp_media}")
        temp_media = None

        # تنظيف الملف بعد الإرسال
        @after_this_request
        def cleanup(response):
            if safe_remove_file(output_file):
                logging.info(f"تم حذف ملف المخرجات: {output_file}")
            return response

        # إرجاع الملف حسب النوع
        if download_type == 'audio':
            download_name = f'clip_{start_time.replace(":", "-")}_{end_time.replace(":", "-")}.mp3'
            mimetype = 'audio/mpeg'
        else:
            download_name = f'clip_{start_time.replace(":", "-")}_{end_time.replace(":", "-")}_{quality}.mp4'
            mimetype = 'video/mp4'

        file_size = os.path.getsize(output_file) if os.path.exists(output_file) else 0
        duration_ms = int((time.time() - start_time_log) * 1000)
        log_activity('video_cutter', 'cut', 'success', duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'quality': quality, 'type': download_type}))

        return send_file(
            output_file,
            as_attachment=True,
            download_name=download_name,
            mimetype=mimetype)

    except subprocess.CalledProcessError as e:
        logging.error(f"خطأ في معالجة الفيديو: {e}")
        safe_remove_files(temp_media, output_file)
        duration_ms = int((time.time() - start_time_log) * 1000)
        error_msg = e.stderr.decode() if e.stderr else str(e)
        log_activity('video_cutter', 'cut', 'error', duration_ms=duration_ms, error_message=error_msg)
        log_error('SubprocessError', error_msg, traceback.format_exc(), tool_name='video_cutter')
        return jsonify({
            'error':
            f'خطأ في معالجة الفيديو: {error_msg}'
        }), 500
    except Exception as e:
        logging.error(f"خطأ عام: {str(e)}")
        safe_remove_files(temp_media, output_file)
        duration_ms = int((time.time() - start_time_log) * 1000)
        log_activity('video_cutter', 'cut', 'error', duration_ms=duration_ms, error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='video_cutter')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


@app.route('/search-anime', methods=['POST'])
def search_anime():
    start_time = time.time()
    temp_path = None
    file_size = 0
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'الرجاء تحميل صورة'}), 400

        file = request.files['image']

        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not allowed_image(file.filename):
            return jsonify({
                'error':
                'نوع الملف غير مدعوم. الأنواع المدعومة: PNG, JPG, JPEG, GIF, WEBP'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        ext = file.filename.rsplit('.', 1)[1].lower()
        temp_path = os.path.join(UPLOAD_FOLDER,
                                 f'anime_search_{unique_id}.{ext}')
        file.save(temp_path)

        from urllib.parse import quote

        gemini_result = identify_anime_with_gemini(temp_path)
        if gemini_result:
            safe_remove_file(temp_path)

            anime_name = gemini_result
            episode_info = 'غير معروف'
            description = ''

            if '|' in gemini_result:
                parts = gemini_result.split('|')
                anime_name = parts[0].strip()
                if len(parts) > 1:
                    episode_info = parts[1].strip() if parts[1].strip().lower(
                    ) != 'unknown' else 'غير معروف'
                if len(parts) > 2:
                    description = parts[2].strip()

            search_links = {
                'anilist':
                f"https://anilist.co/search/anime?search={quote(anime_name)}",
                'myanimelist':
                f"https://myanimelist.net/anime.php?q={quote(anime_name)}",
                'crunchyroll':
                f"https://www.crunchyroll.com/search?q={quote(anime_name)}",
                'youtube':
                f"https://www.youtube.com/results?search_query={quote(anime_name + ' anime')}"
            }

            duration_ms = int((time.time() - start_time) * 1000)
            log_activity('anime_detection_image', 'search', 'success', duration_ms=duration_ms, 
                        file_size=file_size, details=json.dumps({'method': 'gemini_ai', 'anime': anime_name}))

            return jsonify({
                'found': True,
                'anime_name': anime_name,
                'episode': episode_info,
                'similarity': 'AI',
                'timestamp': '',
                'video_preview': '',
                'image_preview': '',
                'detection_method': 'gemini_ai',
                'description': description,
                'search_links': search_links,
                'search_link': search_links['anilist']
            })

        with open(temp_path, 'rb') as f:
            response = requests.post('https://api.trace.moe/search',
                                     files={'image': f},
                                     timeout=30)

        if response.status_code != 200:
            safe_remove_file(temp_path)
            return jsonify({
                'found': False,
                'message':
                'فشل الاتصال بخدمة البحث. تأكد من إعداد مفاتيح Gemini API أو جرب لاحقاً.',
                'suggest_search_by_name': True
            })

        data = response.json()

        if data.get('error'):
            safe_remove_file(temp_path)
            return jsonify({'error': data['error']}), 400

        results = data.get('result', [])

        if not results:
            safe_remove_file(temp_path)
            return jsonify({
                'found': False,
                'message':
                'لم يتم العثور على نتائج. جرب استخدام صورة أوضح أو البحث بالاسم.',
                'suggest_search_by_name': True
            })

        top_result = results[0]
        similarity = top_result.get('similarity', 0)

        if similarity < ANIME_SIMILARITY_THRESHOLD:
            safe_remove_file(temp_path)
            return jsonify({
                'found': False,
                'message':
                'لم يتم العثور على تطابق دقيق. الرجاء استخدام صورة أوضح أو جرب البحث بالاسم.',
                'suggest_search_by_name': True
            })

        safe_remove_file(temp_path)

        anilist_info = top_result.get('anilist', {})
        anime_name = 'غير معروف'

        if isinstance(anilist_info, dict):
            title_info = anilist_info.get('title', {})
            anime_name = title_info.get('english') or title_info.get(
                'romaji') or title_info.get('native') or 'غير معروف'
        elif isinstance(anilist_info, int):
            try:
                anilist_query = '''
                query ($id: Int) {
                    Media (id: $id, type: ANIME) {
                        title {
                            romaji
                            english
                            native
                        }
                    }
                }
                '''
                anilist_response = requests.post('https://graphql.anilist.co',
                                                 json={
                                                     'query': anilist_query,
                                                     'variables': {
                                                         'id': anilist_info
                                                     }
                                                 },
                                                 timeout=10)
                if anilist_response.status_code == 200:
                    anilist_data = anilist_response.json()
                    media = anilist_data.get('data', {}).get('Media', {})
                    title_info = media.get('title', {})
                    anime_name = title_info.get('english') or title_info.get(
                        'romaji') or title_info.get('native') or 'غير معروف'
            except Exception:
                anime_name = f'Anilist ID: {anilist_info}'

        episode = top_result.get('episode', 'غير معروف')
        from_time = top_result.get('from', 0)
        to_time = top_result.get('to', 0)

        def format_time(seconds):
            mins = int(seconds) // 60
            secs = int(seconds) % 60
            return f"{mins:02d}:{secs:02d}"

        result = {
            'found': True,
            'anime_name': anime_name,
            'episode': episode if episode else 'غير معروف',
            'similarity': round(similarity * 100, 2),
            'timestamp': f"{format_time(from_time)} - {format_time(to_time)}",
            'video_preview': top_result.get('video', ''),
            'image_preview': top_result.get('image', '')
        }

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_image', 'search', 'success', duration_ms=duration_ms, 
                    file_size=file_size, details=json.dumps({'method': 'trace.moe', 'anime': anime_name}))

        return jsonify(result)

    except requests.Timeout:
        safe_remove_file(temp_path)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_image', 'search', 'error', duration_ms=duration_ms, 
                    file_size=file_size, error_message='Connection timeout')
        return jsonify(
            {'error': 'انتهت مهلة الاتصال. الرجاء المحاولة مرة أخرى.'}), 500
    except Exception as e:
        safe_remove_file(temp_path)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_image', 'search', 'error', duration_ms=duration_ms, 
                    file_size=file_size, error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='anime_detection_image')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


@app.route('/search-anime-by-name', methods=['POST'])
def search_anime_by_name():
    start_time = time.time()
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        is_description = data.get('is_description', True)

        if not name:
            return jsonify({'error': 'الرجاء إدخال وصف الأنمي أو اسمه'}), 400

        search_term = name
        gemini_suggestion = None

        gemini_result = identify_anime_by_description(name)
        if gemini_result and gemini_result.upper() != 'UNKNOWN':
            if '|' in gemini_result:
                parts = gemini_result.split('|')
                search_term = parts[0].strip()
                confidence = parts[1].strip() if len(parts) > 1 else 'Medium'
                alternatives = parts[2].strip() if len(parts) > 2 else ''
                gemini_suggestion = {
                    'name':
                    search_term,
                    'confidence':
                    confidence,
                    'alternatives':
                    [a.strip() for a in alternatives.split(',') if a.strip()],
                    'original_input':
                    name
                }
            else:
                search_term = gemini_result.strip()
                gemini_suggestion = {
                    'name': search_term,
                    'confidence': 'Medium',
                    'alternatives': [],
                    'original_input': name
                }

        anilist_query = '''
        query ($search: String) {
            Page(page: 1, perPage: 5) {
                media(search: $search, type: ANIME, sort: POPULARITY_DESC) {
                    id
                    title {
                        romaji
                        english
                        native
                    }
                    description(asHtml: false)
                    coverImage {
                        large
                        medium
                    }
                    episodes
                    status
                    genres
                    averageScore
                    seasonYear
                    siteUrl
                }
            }
        }
        '''

        response = requests.post('https://graphql.anilist.co',
                                 json={
                                     'query': anilist_query,
                                     'variables': {
                                         'search': search_term
                                     }
                                 },
                                 timeout=15)

        if response.status_code != 200:
            return jsonify({'error': 'فشل الاتصال بخدمة البحث'}), 500

        data = response.json()
        media_list = data.get('data', {}).get('Page', {}).get('media', [])

        if not media_list:
            return jsonify({
                'found': False,
                'message':
                'لم يتم العثور على نتائج. جرب اسماً أو وصفاً مختلفاً.',
                'gemini_suggestion': gemini_suggestion
            })

        from urllib.parse import quote
        results = []
        for media in media_list:
            title_info = media.get('title', {})
            anime_name = title_info.get('english') or title_info.get(
                'romaji') or title_info.get('native') or 'غير معروف'

            description = media.get('description', '')
            if description:
                description = re.sub(r'<[^>]+>', '', description)
                if len(description) > 300:
                    description = description[:300] + '...'

            results.append({
                'id':
                media.get('id'),
                'name':
                anime_name,
                'name_native':
                title_info.get('native', ''),
                'description':
                description or 'لا يوجد وصف متاح',
                'cover':
                media.get('coverImage', {}).get('large')
                or media.get('coverImage', {}).get('medium', ''),
                'episodes':
                media.get('episodes') or 'غير محدد',
                'status':
                media.get('status', 'غير معروف'),
                'genres':
                media.get('genres', []),
                'score':
                media.get('averageScore') or 0,
                'year':
                media.get('seasonYear') or 'غير معروف',
                'anilist_url':
                media.get('siteUrl', ''),
                'search_links': {
                    'anilist':
                    media.get(
                        'siteUrl',
                        f"https://anilist.co/search/anime?search={quote(anime_name)}"
                    ),
                    'myanimelist':
                    f"https://myanimelist.net/anime.php?q={quote(anime_name)}",
                    'crunchyroll':
                    f"https://www.crunchyroll.com/search?q={quote(anime_name)}"
                }
            })

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_name', 'search', 'success', duration_ms=duration_ms, 
                    details=json.dumps({'search_term': search_term, 'results_count': len(results)}))

        return jsonify({
            'found': True,
            'results': results,
            'gemini_suggestion': gemini_suggestion
        })

    except requests.Timeout:
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_name', 'search', 'error', duration_ms=duration_ms, 
                    error_message='Connection timeout')
        return jsonify(
            {'error': 'انتهت مهلة الاتصال. الرجاء المحاولة مرة أخرى.'}), 500
    except Exception as e:
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('anime_detection_name', 'search', 'error', duration_ms=duration_ms, 
                    error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='anime_detection_name')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


@app.route('/transcribe-file', methods=['POST'])
def transcribe_file():
    start_time = time.time()
    temp_audio = None
    temp_wav = None
    chunk_files = []
    file_size = 0

    try:
        if 'audio' not in request.files:
            log_activity('audio_transcription', 'transcribe', 'error', error_message='No audio file uploaded')
            return jsonify({'error': 'الرجاء تحميل ملف صوتي'}), 400

        file = request.files['audio']
        language = request.form.get('language', 'ar')
        if language not in LANGUAGE_NAMES:
            language = 'ar'

        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not allowed_audio(file.filename):
            return jsonify({
                'error':
                'نوع الملف غير مدعوم. الأنواع المدعومة: MP3, WAV, OGG, M4A, FLAC, AAC, WMA'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        ext = file.filename.rsplit('.', 1)[1].lower()
        temp_audio = os.path.join(UPLOAD_FOLDER, f'audio_{unique_id}.{ext}')
        temp_wav = os.path.join(UPLOAD_FOLDER, f'audio_{unique_id}.wav')

        file.save(temp_audio)
        file_size = os.path.getsize(temp_audio) if os.path.exists(temp_audio) else 0

        try:
            audio = AudioSegment.from_file(temp_audio)
            audio = audio.set_channels(1)
            audio = audio.set_frame_rate(16000)
            duration_seconds = len(audio) / 1000
            duration_minutes = duration_seconds / 60
            logging.info(
                f"Audio file loaded: {duration_minutes:.1f} minutes ({duration_seconds:.0f} seconds)"
            )
        except Exception as e:
            logging.error(f"Audio loading error: {str(e)}")
            duration_ms = int((time.time() - start_time) * 1000)
            log_activity('audio_transcription', 'transcribe', 'error', duration_ms=duration_ms, 
                        file_size=file_size, error_message=f'Audio loading error: {str(e)}')
            return jsonify({'error':
                            f'خطأ في قراءة الملف الصوتي: {str(e)}'}), 400

        if duration_seconds > 7200:
            safe_remove_file(temp_audio)
            return jsonify({
                'error':
                'الملف الصوتي أطول من ساعتين. الحد الأقصى المدعوم هو ساعتين.'
            }), 400

        text = None
        transcription_method = 'gemini'

        if GEMINI_KEYS:
            logging.info(
                f"Audio duration: {duration_minutes:.1f} min - using Gemini AI with {get_available_keys_count()} available keys"
            )

            try:
                text = transcribe_audio_with_gemini(temp_audio, language)
            except ValueError as ve:
                safe_remove_file(temp_audio)
                run_garbage_collection()
                return jsonify({'error': str(ve)}), 400
            except Exception as e:
                logging.error(f"Gemini transcription failed: {str(e)}")
                text = None

            if text:
                safe_remove_file(temp_audio)
                run_garbage_collection()
                duration_ms = int((time.time() - start_time) * 1000)
                log_activity('audio_transcription', 'transcribe', 'success', 
                            duration_ms=duration_ms, file_size=file_size,
                            details=json.dumps({'method': 'gemini_ai', 'audio_duration': round(duration_seconds, 1)}))

                return jsonify({
                    'success': True,
                    'text': text,
                    'duration': round(duration_seconds, 1),
                    'duration_formatted': f"{int(duration_minutes)} دقيقة",
                    'method': 'gemini_ai',
                    'keys_available': get_available_keys_count()
                })

        if not text and not GEMINI_KEYS:
            transcription_method = 'google_speech'
            logging.info(
                f"No Gemini keys available, falling back to Google Speech Recognition for {duration_seconds}s audio"
            )

            if duration_seconds > 180:
                safe_remove_file(temp_audio)
                return jsonify({
                    'error':
                    'لتحويل ملفات أطول من 3 دقائق، يرجى إضافة مفاتيح Gemini API. الملفات الطويلة تحتاج إلى مفاتيح API.'
                }), 400

            try:
                audio.export(temp_wav, format="wav")
            except Exception as e:
                logging.error(f"Audio conversion error: {str(e)}")
                return jsonify(
                    {'error': f'خطأ في تحويل الملف الصوتي: {str(e)}'}), 400

            recognizer = sr.Recognizer()
            audio_segment = AudioSegment.from_wav(temp_wav)

            if duration_seconds > 180:
                audio_segment = audio_segment[:180 * 1000]
                duration_seconds = 180
                logging.info("Audio trimmed to 180 seconds for Google Speech")

            CHUNK_DURATION_MS = 30 * 1000

            if duration_seconds <= 30:
                with sr.AudioFile(temp_wav) as source:
                    audio_data = recognizer.record(source)

                try:
                    text = recognizer.recognize_google(audio_data,
                                                       language='ar-EG')
                except sr.UnknownValueError:
                    try:
                        text = recognizer.recognize_google(audio_data,
                                                           language='ar-SA')
                    except sr.UnknownValueError:
                        text = ''
                except sr.RequestError as e:
                    return jsonify(
                        {'error':
                         f'خطأ في خدمة التعرف على الصوت: {str(e)}'}), 500
            else:
                logging.info(f"Splitting audio into chunks for Google Speech")

                transcribed_texts = []
                num_chunks = int(len(audio_segment) / CHUNK_DURATION_MS) + 1

                for i in range(num_chunks):
                    start_ms = i * CHUNK_DURATION_MS
                    end_ms = min((i + 1) * CHUNK_DURATION_MS,
                                 len(audio_segment))
                    chunk = audio_segment[start_ms:end_ms]

                    chunk_path = os.path.join(UPLOAD_FOLDER,
                                              f'chunk_{unique_id}_{i}.wav')
                    chunk_files.append(chunk_path)
                    chunk.export(chunk_path, format="wav")

                    try:
                        with sr.AudioFile(chunk_path) as source:
                            chunk_audio = recognizer.record(source)

                        chunk_text = recognizer.recognize_google(
                            chunk_audio, language='ar-EG')
                        transcribed_texts.append(chunk_text)
                        logging.info(
                            f"Chunk {i+1}/{num_chunks} transcribed successfully"
                        )
                    except sr.UnknownValueError:
                        try:
                            if chunk_audio:
                                chunk_text = recognizer.recognize_google(
                                    chunk_audio, language='ar-SA')
                                transcribed_texts.append(chunk_text)
                            logging.info(
                                f"Chunk {i+1}/{num_chunks} transcribed with fallback"
                            )
                        except sr.UnknownValueError:
                            logging.info(
                                f"Chunk {i+1}/{num_chunks}: No speech detected"
                            )
                            continue
                    except sr.RequestError as e:
                        logging.error(
                            f"Chunk {i+1}/{num_chunks} error: {str(e)}")
                        continue

                text = ' '.join(transcribed_texts)

        safe_remove_files(temp_audio, temp_wav, *chunk_files)
        run_garbage_collection()
        duration_ms = int((time.time() - start_time) * 1000)

        if not text:
            log_activity('audio_transcription', 'transcribe', 'warning', 
                        duration_ms=duration_ms, file_size=file_size,
                        details=json.dumps({'method': transcription_method, 'result': 'no_speech_detected'}))
            return jsonify({
                'success':
                True,
                'text':
                '',
                'message':
                'لم يتم التعرف على أي كلام في الملف الصوتي. تأكد من إعداد مفاتيح Gemini API للملفات الطويلة.'
            })

        log_activity('audio_transcription', 'transcribe', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'method': transcription_method, 'audio_duration': round(duration_seconds, 1)}))
        return jsonify({
            'success': True,
            'text': text,
            'duration': round(duration_seconds, 1),
            'method': transcription_method
        })

    except ValueError as ve:
        logging.error(f"Transcription ValueError: {str(ve)}")
        safe_remove_files(temp_audio, temp_wav, *chunk_files)
        run_garbage_collection()
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('audio_transcription', 'transcribe', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(ve))
        return jsonify({'error': str(ve)}), 400
    except Exception as e:
        error_message = str(e)
        logging.error(f"Transcription error: {error_message}")
        safe_remove_files(temp_audio, temp_wav, *chunk_files)
        run_garbage_collection()
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('audio_transcription', 'transcribe', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=error_message[:200])
        log_error('TranscriptionError', error_message, traceback.format_exc(), tool_name='audio_transcription')

        if 'timeout' in error_message.lower(
        ) or 'timed out' in error_message.lower():
            return jsonify({
                'error':
                'انتهت مهلة المعالجة. جاري المحاولة مع مفتاح آخر أو حاول مرة أخرى.'
            }), 500
        elif 'quota' in error_message.lower() or 'rate' in error_message.lower(
        ):
            return jsonify({
                'error':
                'تم استنفاد حصة API. حاول مرة أخرى بعد ساعة أو أضف مفاتيح جديدة.'
            }), 500
        elif 'memory' in error_message.lower():
            return jsonify({'error':
                            'الملف كبير جداً. جرب ملف أصغر حجماً.'}), 500
        return jsonify({'error':
                        f'خطأ في التحويل: {error_message[:100]}'}), 500


@app.route('/ocr-image', methods=['POST'])
def ocr_image():
    start_time = time.time()
    file_size = 0
    temp_image = None

    try:
        if 'image' not in request.files:
            return jsonify({'error': 'الرجاء تحميل صورة'}), 400

        file = request.files['image']

        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not allowed_image(file.filename):
            return jsonify({
                'error':
                'نوع الملف غير مدعوم. الأنواع المدعومة: PNG, JPG, JPEG, GIF, WEBP'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        ext = file.filename.rsplit('.', 1)[1].lower()
        temp_image = os.path.join(UPLOAD_FOLDER, f'ocr_{unique_id}.{ext}')

        file.save(temp_image)
        file_size = os.path.getsize(temp_image) if os.path.exists(temp_image) else 0

        image = Image.open(temp_image)

        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(image,
                                           lang='ara+eng',
                                           config=custom_config)

        safe_remove_file(temp_image)

        duration_ms = int((time.time() - start_time) * 1000)
        if not text.strip():
            log_activity('ocr_extraction', 'extract', 'success', 
                        duration_ms=duration_ms, file_size=file_size,
                        details=json.dumps({'text_length': 0}))
            return jsonify({
                'success': True,
                'text': '',
                'message': 'لم يتم العثور على نص في الصورة'
            })

        log_activity('ocr_extraction', 'extract', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'text_length': len(text.strip())}))
        return jsonify({'success': True, 'text': text.strip()})

    except Exception as e:
        safe_remove_file(temp_image)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('ocr_extraction', 'extract', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e))
        log_error('OCRError', str(e), traceback.format_exc(), tool_name='ocr_extraction')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


@app.route('/search-podcast-by-image', methods=['POST'])
def search_podcast_by_image():
    start_time = time.time()
    file_size = 0
    temp_image = None

    try:
        if 'image' not in request.files:
            return jsonify({'error': 'الرجاء تحميل صورة'}), 400

        file = request.files['image']

        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not allowed_image(file.filename):
            return jsonify({
                'error':
                'نوع الملف غير مدعوم. الأنواع المدعومة: PNG, JPG, JPEG, GIF, WEBP'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        ext = file.filename.rsplit('.', 1)[1].lower()
        temp_image = os.path.join(UPLOAD_FOLDER,
                                  f'podcast_ocr_{unique_id}.{ext}')

        file.save(temp_image)
        file_size = os.path.getsize(temp_image) if os.path.exists(temp_image) else 0

        search_term = None
        extracted_text = ''
        detection_method = 'ai_vision'
        host_names = ''
        platform_hint = ''

        gemini_result = identify_podcast_with_gemini(temp_image)
        if gemini_result:
            if '|' in gemini_result:
                parts = gemini_result.split('|')
                search_term = parts[0].strip()
                extracted_text = search_term
                if len(parts) > 1:
                    host_names = parts[1].strip() if parts[1].strip().lower(
                    ) != 'unknown' else ''
                if len(parts) > 2:
                    platform_hint = parts[2].strip()
            else:
                search_term = gemini_result.strip()
                extracted_text = search_term
            detection_method = 'gemini_ai'

        if not search_term:
            image = Image.open(temp_image)
            extracted_text = pytesseract.image_to_string(
                image, lang='ara+eng', config='--oem 3 --psm 6')
            extracted_text = extracted_text.strip()
            detection_method = 'ocr'

            if extracted_text and len(extracted_text) > 3:
                search_term = ' '.join(extracted_text.split()[:10])

        if not search_term:
            safe_remove_file(temp_image)
            return jsonify({
                'found':
                False,
                'message':
                'لم يتم التعرف على البودكاست. جرب صورة أوضح تحتوي على نص أو وجوه معروفة.'
            })

        itunes_response = requests.get('https://itunes.apple.com/search',
                                       params={
                                           'term': search_term,
                                           'media': 'podcast',
                                           'limit': 5
                                       },
                                       timeout=15)

        if itunes_response.status_code != 200:
            safe_remove_file(temp_image)
            return jsonify({'error': 'فشل الاتصال بخدمة البحث'}), 500

        itunes_data = itunes_response.json()
        podcasts = itunes_data.get('results', [])

        smart_links = generate_podcast_search_links(search_term)

        if not podcasts:
            safe_remove_file(temp_image)
            return jsonify({
                'found': False,
                'extracted_text': extracted_text,
                'detection_method': detection_method,
                'message':
                'لم يتم العثور على بودكاست مطابق في iTunes. جرب الروابط أدناه للبحث يدوياً.',
                'search_links': smart_links
            })

        best_match = None
        best_similarity = 0

        for podcast in podcasts:
            podcast_name = podcast.get('collectionName', '')
            artist_name = podcast.get('artistName', '')

            name_similarity = fuzz.ratio(extracted_text.lower(),
                                         podcast_name.lower())
            artist_similarity = fuzz.ratio(extracted_text.lower(),
                                           artist_name.lower())

            partial_name = fuzz.partial_ratio(extracted_text.lower(),
                                              podcast_name.lower())

            max_similarity = max(name_similarity, artist_similarity,
                                 partial_name)

            if max_similarity > best_similarity:
                best_similarity = max_similarity
                best_match = podcast

        if best_similarity < PODCAST_NAME_SIMILARITY_THRESHOLD and detection_method == 'ocr':
            gemini_result = identify_podcast_with_gemini(temp_image)

            if gemini_result:
                detection_method = 'ai_vision'
                extracted_text = gemini_result

                itunes_response2 = requests.get(
                    'https://itunes.apple.com/search',
                    params={
                        'term': gemini_result,
                        'media': 'podcast',
                        'limit': 5
                    },
                    timeout=15)

                if itunes_response2.status_code == 200:
                    itunes_data2 = itunes_response2.json()
                    podcasts2 = itunes_data2.get('results', [])

                    if podcasts2:
                        best_match = None
                        best_similarity = 0

                        for podcast in podcasts2:
                            podcast_name = podcast.get('collectionName', '')
                            artist_name = podcast.get('artistName', '')

                            name_similarity = fuzz.ratio(
                                gemini_result.lower(), podcast_name.lower())
                            artist_similarity = fuzz.ratio(
                                gemini_result.lower(), artist_name.lower())
                            partial_name = fuzz.partial_ratio(
                                gemini_result.lower(), podcast_name.lower())

                            max_similarity = max(name_similarity,
                                                 artist_similarity,
                                                 partial_name)

                            if max_similarity > best_similarity:
                                best_similarity = max_similarity
                                best_match = podcast

                        smart_links = generate_podcast_search_links(
                            gemini_result)

            if not best_match or best_similarity < PODCAST_NAME_SIMILARITY_THRESHOLD:
                safe_remove_file(temp_image)
                return jsonify({
                    'found': False,
                    'extracted_text': extracted_text,
                    'detection_method': detection_method,
                    'similarity': best_similarity,
                    'message':
                    'تعذر التعرف على البودكاست بدقة عالية. جرب الروابط أدناه للبحث يدوياً.',
                    'search_links': smart_links
                })

        safe_remove_file(temp_image)

        if not best_match:
            return jsonify({
                'found': False,
                'message': 'لم يتم العثور على بودكاست مطابق',
                'search_links': smart_links
            })

        podcast_name = best_match.get('collectionName', 'غير معروف')
        final_smart_links = generate_podcast_search_links(podcast_name)

        result = {
            'found': True,
            'extracted_text': extracted_text,
            'detection_method': detection_method,
            'similarity': best_similarity,
            'host_names': host_names,
            'platform_hint': platform_hint,
            'podcast': {
                'name':
                podcast_name,
                'artist':
                best_match.get('artistName', 'غير معروف'),
                'artwork':
                best_match.get('artworkUrl600')
                or best_match.get('artworkUrl100', ''),
                'genre':
                best_match.get('primaryGenreName', 'غير محدد'),
                'episode_count':
                best_match.get('trackCount', 0),
                'feed_url':
                best_match.get('feedUrl', ''),
                'itunes_url':
                best_match.get('collectionViewUrl', '')
            },
            'search_links': final_smart_links
        }

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_image', 'detect', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'podcast': podcast_name, 'method': detection_method}))
        return jsonify(result)

    except requests.Timeout:
        safe_remove_file(temp_image)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_image', 'detect', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message='Connection timeout')
        return jsonify(
            {'error': 'انتهت مهلة الاتصال. الرجاء المحاولة مرة أخرى.'}), 500
    except Exception as e:
        safe_remove_file(temp_image)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_image', 'detect', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e))
        log_error('PodcastImageError', str(e), traceback.format_exc(), tool_name='podcast_detection_image')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


@app.route('/search-podcast-by-audio', methods=['POST'])
def search_podcast_by_audio():
    start_time = time.time()
    file_size = 0
    temp_audio = None
    temp_wav = None

    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'الرجاء تحميل ملف صوتي'}), 400

        file = request.files['audio']

        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not allowed_audio(file.filename):
            return jsonify({
                'error':
                'نوع الملف غير مدعوم. الأنواع المدعومة: MP3, WAV, OGG, M4A, FLAC, AAC, WMA'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        ext = file.filename.rsplit('.', 1)[1].lower()
        temp_audio = os.path.join(UPLOAD_FOLDER,
                                  f'podcast_audio_{unique_id}.{ext}')
        temp_wav = os.path.join(UPLOAD_FOLDER,
                                f'podcast_audio_{unique_id}.wav')

        file.save(temp_audio)
        file_size = os.path.getsize(temp_audio) if os.path.exists(temp_audio) else 0

        try:
            audio = AudioSegment.from_file(temp_audio)
            audio = audio.set_channels(1)
            audio = audio.set_frame_rate(16000)

            if len(audio) > 60000:
                audio = audio[:60000]

            audio.export(temp_wav, format="wav")
        except Exception as e:
            return jsonify({'error':
                            f'خطأ في تحويل الملف الصوتي: {str(e)}'}), 400

        recognizer = sr.Recognizer()

        with sr.AudioFile(temp_wav) as source:
            audio_data = recognizer.record(source)

        transcribed_text = ''
        try:
            transcribed_text = recognizer.recognize_google(audio_data,
                                                           language='ar-SA')
        except sr.UnknownValueError:
            try:
                transcribed_text = recognizer.recognize_google(
                    audio_data, language='en-US')
            except sr.UnknownValueError:
                transcribed_text = ''
        except sr.RequestError as e:
            return jsonify(
                {'error': f'خطأ في خدمة التعرف على الصوت: {str(e)}'}), 500

        safe_remove_files(temp_audio, temp_wav)

        if not transcribed_text:
            return jsonify({
                'found':
                False,
                'message':
                'لم يتم التعرف على أي كلام في الملف الصوتي. جرب ملفاً أوضح.'
            })

        gemini_podcast_name = identify_podcast_from_transcript(
            transcribed_text)

        search_term = gemini_podcast_name if gemini_podcast_name else ' '.join(
            transcribed_text.split()[:15])
        detection_method = 'gemini_ai' if gemini_podcast_name else 'keyword_match'

        itunes_response = requests.get('https://itunes.apple.com/search',
                                       params={
                                           'term': search_term,
                                           'media': 'podcast',
                                           'limit': 5
                                       },
                                       timeout=15)

        smart_links = generate_podcast_search_links(search_term)

        if itunes_response.status_code != 200:
            return jsonify({
                'found': False,
                'transcribed_text': transcribed_text,
                'identified_name': gemini_podcast_name,
                'detection_method': detection_method,
                'message':
                'فشل الاتصال بخدمة iTunes. استخدم الروابط أدناه للبحث يدوياً.',
                'search_links': smart_links
            })

        itunes_data = itunes_response.json()
        podcasts = itunes_data.get('results', [])

        if not podcasts:
            return jsonify({
                'found': False,
                'transcribed_text': transcribed_text,
                'identified_name': gemini_podcast_name,
                'detection_method': detection_method,
                'message':
                'لم يتم العثور على بودكاست مطابق. استخدم الروابط أدناه للبحث يدوياً.',
                'search_links': smart_links
            })

        results = []
        for podcast in podcasts:
            podcast_name = podcast.get('collectionName', 'غير معروف')
            podcast_links = generate_podcast_search_links(podcast_name)
            results.append({
                'name':
                podcast_name,
                'artist':
                podcast.get('artistName', 'غير معروف'),
                'artwork':
                podcast.get('artworkUrl600')
                or podcast.get('artworkUrl100', ''),
                'genre':
                podcast.get('primaryGenreName', 'غير محدد'),
                'episode_count':
                podcast.get('trackCount', 0),
                'feed_url':
                podcast.get('feedUrl', ''),
                'itunes_url':
                podcast.get('collectionViewUrl', ''),
                'search_links':
                podcast_links
            })

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_audio', 'detect', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'found': True, 'method': detection_method}))
        return jsonify({
            'found': True,
            'transcribed_text': transcribed_text,
            'identified_name': gemini_podcast_name,
            'detection_method': detection_method,
            'results': results,
            'search_links': smart_links
        })

    except requests.Timeout:
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_audio', 'detect', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message='Connection timeout')
        return jsonify(
            {'error': 'انتهت مهلة الاتصال. الرجاء المحاولة مرة أخرى.'}), 500
    except Exception as e:
        safe_remove_files(temp_audio, temp_wav)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('podcast_detection_audio', 'detect', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e))
        log_error('PodcastAudioError', str(e), traceback.format_exc(), tool_name='podcast_detection_audio')
        return jsonify({'error': f'خطأ: {str(e)}'}), 500


def cleanup_download_files(output_template, output_file):
    """Clean up all possible temp files from download"""
    for ext in [
            'mp4', 'mkv', 'webm', 'avi', 'mov', 'mp3', 'wav', 'ogg', 'm4a',
            'part', 'ytdl', 'temp'
    ]:
        temp_file = f'{output_template}.{ext}'
        safe_remove_file(temp_file)
    safe_remove_file(output_file)


@app.route('/download-video', methods=['POST'])
def download_video():
    start_time = time.time()
    file_size = 0
    output_file = None
    output_template = None
    media_title = 'media'

    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        download_format = data.get('format', 'video')

        if not url:
            return jsonify({'error': 'الرجاء إدخال رابط الفيديو'}), 400

        if not url.startswith(('http://', 'https://')):
            return jsonify({
                'error':
                'الرابط غير صالح. يجب أن يبدأ بـ http:// أو https://'
            }), 400

        unique_id = str(uuid.uuid4())[:8]
        output_template = os.path.join(UPLOAD_FOLDER, f'download_{unique_id}')

        if download_format == 'audio':
            output_file = f'{output_template}.mp3'
        else:
            output_file = f'{output_template}.mp4'

        info_opts = {
            'quiet': True,
            'no_warnings': True,
            'socket_timeout': 1800,
            'noplaylist': True,
            'retries': 3,
            'http_headers': {
                'User-Agent':
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept':
                'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
            },
        }

        if os.path.exists('cookies.txt'):
            info_opts['cookiefile'] = 'cookies.txt'

        with yt_dlp.YoutubeDL(info_opts) as ydl:
            info = ydl.extract_info(url, download=False)

            if not info:
                return jsonify(
                    {'error': 'لم يتم العثور على محتوى في هذا الرابط'}), 400

            if info.get('_type') == 'playlist' or 'entries' in info:
                return jsonify({
                    'error':
                    'لا يمكن تحميل قوائم التشغيل. الرجاء استخدام رابط واحد.'
                }), 400

            media_title = info.get('title', 'media')
            media_title = re.sub(r'[\\/*?:"<>|]', '', media_title)[:50]

            if info.get('is_live'):
                return jsonify({'error': 'لا يمكن تحميل البث المباشر.'}), 400

        is_tiktok = 'tiktok.com' in url.lower() or 'douyin.com' in url.lower()
        is_instagram = 'instagram.com' in url.lower()
        is_twitter = 'twitter.com' in url.lower() or 'x.com' in url.lower()

        common_headers = {
            'User-Agent':
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept':
            'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Sec-Fetch-Mode': 'navigate',
        }

        cookiefile_opt = 'cookies.txt' if os.path.exists(
            'cookies.txt') else None

        if download_format == 'audio':
            ydl_opts = {
                'format':
                'bestaudio[ext=m4a]/bestaudio[ext=mp3]/bestaudio/best',
                'outtmpl':
                output_template + '.%(ext)s',
                'noplaylist':
                True,
                'quiet':
                True,
                'no_warnings':
                True,
                'socket_timeout':
                7200,
                'retries':
                5,
                'http_headers':
                common_headers,
                'postprocessors': [{
                    'key': 'FFmpegExtractAudio',
                    'preferredcodec': 'mp3',
                    'preferredquality': '128',
                }],
            }
            if cookiefile_opt:
                ydl_opts['cookiefile'] = cookiefile_opt
        else:
            if is_tiktok or is_instagram or is_twitter:
                ydl_opts = {
                    'format':
                    'bestvideo[ext=mp4][vcodec^=avc1]+bestaudio/bestvideo[ext=mp4]+bestaudio/best[ext=mp4][vcodec^=avc1]/best[ext=mp4]/best',
                    'merge_output_format':
                    'mp4',
                    'outtmpl':
                    output_template + '.%(ext)s',
                    'noplaylist':
                    True,
                    'quiet':
                    True,
                    'no_warnings':
                    True,
                    'socket_timeout':
                    7200,
                    'retries':
                    10,
                    'fragment_retries':
                    10,
                    'http_headers':
                    common_headers,
                    'extractor_args': {
                        'tiktok': {
                            'api_hostname':
                            'api22-normal-c-useast1a.tiktokv.com'
                        }
                    },
                    'postprocessors': [{
                        'key': 'FFmpegVideoConvertor',
                        'preferedformat': 'mp4',
                    }],
                    'postprocessor_args': {
                        'FFmpegVideoConvertor': [
                            '-c:v', 'libx264', '-preset', 'fast', '-crf', '23',
                            '-c:a', 'aac', '-b:a', '128k', '-movflags',
                            '+faststart'
                        ]
                    },
                }
                if cookiefile_opt:
                    ydl_opts['cookiefile'] = cookiefile_opt
            else:
                ydl_opts = {
                    'format':
                    'bestvideo[ext=mp4][vcodec^=avc1]+bestaudio[ext=m4a]/bestvideo[ext=mp4]+bestaudio[ext=m4a]/bestvideo+bestaudio/best[ext=mp4]/best',
                    'merge_output_format':
                    'mp4',
                    'outtmpl':
                    output_template + '.%(ext)s',
                    'noplaylist':
                    True,
                    'quiet':
                    True,
                    'no_warnings':
                    True,
                    'socket_timeout':
                    7200,
                    'retries':
                    5,
                    'fragment_retries':
                    5,
                    'http_headers':
                    common_headers,
                    'postprocessors': [{
                        'key': 'FFmpegVideoConvertor',
                        'preferedformat': 'mp4',
                    }],
                    'postprocessor_args': {
                        'FFmpegVideoConvertor': [
                            '-c:v', 'libx264', '-c:a', 'aac', '-movflags',
                            '+faststart'
                        ]
                    },
                }
                if cookiefile_opt:
                    ydl_opts['cookiefile'] = cookiefile_opt

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])

        if download_format == 'audio':
            if not os.path.exists(output_file):
                for ext in ['mp3', 'wav', 'm4a', 'ogg', 'opus', 'webm']:
                    alt_file = f'{output_template}.{ext}'
                    if os.path.exists(alt_file):
                        if ext != 'mp3':
                            try:
                                ffmpeg_cmd = [
                                    'ffmpeg', '-y', '-i', alt_file, '-vn',
                                    '-acodec', 'libmp3lame', '-ab', '128k',
                                    '-threads', '4', '-q:a', '5', output_file
                                ]
                                subprocess.run(ffmpeg_cmd,
                                               check=True,
                                               capture_output=True,
                                               timeout=7200)
                                safe_remove_file(alt_file)
                            except subprocess.CalledProcessError:
                                cleanup_download_files(output_template,
                                                       output_file)
                                return jsonify({
                                    'error':
                                    'فشل في تحويل الصوت إلى MP3. جرب رابطاً آخر.'
                                }), 500
                            except subprocess.TimeoutExpired:
                                cleanup_download_files(output_template,
                                                       output_file)
                                return jsonify({
                                    'error':
                                    'انتهت مهلة تحويل الصوت. جرب محتوى أقصر.'
                                }), 500
                        else:
                            output_file = alt_file
                        break
        else:
            if not os.path.exists(output_file):
                for ext in ['mp4', 'mkv', 'webm', 'avi', 'mov']:
                    alt_file = f'{output_template}.{ext}'
                    if os.path.exists(alt_file):
                        needs_conversion = False
                        if ext != 'mp4':
                            needs_conversion = True
                        else:
                            try:
                                probe_cmd = [
                                    'ffprobe', '-v', 'error',
                                    '-select_streams', 'v:0', '-show_entries',
                                    'stream=codec_name', '-of',
                                    'default=noprint_wrappers=1:nokey=1',
                                    alt_file
                                ]
                                result = subprocess.run(probe_cmd,
                                                        capture_output=True,
                                                        text=True,
                                                        timeout=30)
                                codec = result.stdout.strip().lower()
                                if codec in [
                                        'hevc', 'h265', 'hev1', 'hvc1', 'vp9',
                                        'av1'
                                ]:
                                    needs_conversion = True
                                    logging.info(
                                        f"Detected {codec} codec, converting to H.264 for compatibility"
                                    )
                            except Exception as probe_error:
                                logging.warning(
                                    f"Could not probe video codec: {probe_error}"
                                )
                                needs_conversion = True

                        if needs_conversion:
                            try:
                                ffmpeg_cmd = [
                                    'ffmpeg', '-y', '-i', alt_file, '-c:v',
                                    'libx264', '-preset', 'fast', '-crf', '23',
                                    '-c:a', 'aac', '-b:a', '128k', '-movflags',
                                    '+faststart', '-strict', 'experimental',
                                    output_file
                                ]
                                subprocess.run(ffmpeg_cmd,
                                               check=True,
                                               capture_output=True,
                                               timeout=7200)
                                if alt_file != output_file:
                                    safe_remove_file(alt_file)
                            except subprocess.CalledProcessError:
                                cleanup_download_files(output_template,
                                                       output_file)
                                return jsonify({
                                    'error':
                                    'فشل في تحويل الفيديو إلى MP4. جرب رابطاً آخر.'
                                }), 500
                            except subprocess.TimeoutExpired:
                                cleanup_download_files(output_template,
                                                       output_file)
                                return jsonify({
                                    'error':
                                    'انتهت مهلة تحويل الفيديو. جرب فيديو أقصر.'
                                }), 500
                        else:
                            output_file = alt_file
                        break

        if not os.path.exists(output_file):
            cleanup_download_files(output_template, output_file)
            if download_format == 'audio':
                return jsonify(
                    {'error': 'فشل في تحميل الصوت. جرب رابطاً آخر.'}), 500
            else:
                return jsonify(
                    {'error': 'فشل في تحميل الفيديو. جرب رابطاً آخر.'}), 500

        final_output = output_file
        final_template = output_template
        file_size = os.path.getsize(output_file) if os.path.exists(output_file) else 0

        @after_this_request
        def cleanup(response):
            cleanup_download_files(final_template, final_output)
            return response

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('media_download', 'download', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'format': download_format, 'title': media_title[:50]}))
        if download_format == 'audio':
            return send_file(output_file,
                             as_attachment=True,
                             download_name=f'{media_title}.mp3',
                             mimetype='audio/mpeg')
        else:
            return send_file(output_file,
                             as_attachment=True,
                             download_name=f'{media_title}.mp4',
                             mimetype='video/mp4')

    except yt_dlp.utils.DownloadError as e:
        cleanup_download_files(output_template, output_file)
        error_msg = str(e).lower()
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('media_download', 'download', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e)[:200])
        if 'private' in error_msg:
            return jsonify({'error': 'هذا الفيديو خاص ولا يمكن تحميله.'}), 400
        elif 'unavailable' in error_msg or 'removed' in error_msg or 'deleted' in error_msg:
            return jsonify({'error': 'الفيديو غير متاح أو تم حذفه.'}), 400
        elif 'unsupported' in error_msg or 'no video formats' in error_msg:
            return jsonify(
                {'error': 'هذا الموقع غير مدعوم أو الرابط غير صحيح.'}), 400
        elif '403' in error_msg or 'forbidden' in error_msg:
            return jsonify({'error':
                            'تم رفض الوصول. الفيديو قد يكون محمياً.'}), 400
        elif 'geo' in error_msg or 'country' in error_msg:
            return jsonify({'error': 'هذا الفيديو غير متاح في منطقتك.'}), 400
        elif 'drm' in error_msg or 'protected' in error_msg:
            return jsonify(
                {'error':
                 'هذا الفيديو محمي بحقوق النشر ولا يمكن تحميله.'}), 400
        elif 'age' in error_msg or 'sign in' in error_msg:
            return jsonify({'error': 'هذا الفيديو يتطلب تسجيل الدخول.'}), 400
        else:
            return jsonify({'error': f'خطأ في التحميل: {str(e)[:100]}'}), 400
    except Exception as e:
        cleanup_download_files(output_template, output_file)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('media_download', 'download', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e)[:200])
        log_error('DownloadError', str(e), traceback.format_exc(), tool_name='media_download')
        return jsonify({'error': f'خطأ غير متوقع: {str(e)[:100]}'}), 500


def extract_arabic_text_with_tesseract(image):
    """Use Tesseract OCR to extract Arabic text from PDF page image (no Gemini API usage)"""
    try:
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(image, lang='ara+eng', config=custom_config)
        if text and text.strip():
            return text.strip()
    except Exception as e:
        logging.warning(f"Tesseract OCR failed: {e}")
    return None


def convert_pdf_with_ocr(pdf_path, docx_path, use_tesseract=True):
    """Convert PDF to DOCX using Tesseract OCR for Arabic text extraction (no Gemini API usage)"""
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    try:
        pdf_doc = fitz.open(pdf_path)
        doc = Document()

        style = doc.styles['Normal']
        style.font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        total_pages = len(pdf_doc)
        extracted_any_text = False

        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            extracted_text = None

            if use_tesseract:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                extracted_text = extract_arabic_text_with_tesseract(img)

            if not extracted_text:
                extracted_text = page.get_text("text")

            if extracted_text and extracted_text.strip():
                extracted_any_text = True
                paragraphs = extracted_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run = p.add_run(para_text.strip())
                        run.font.size = Pt(12)

            if page_num < total_pages - 1:
                doc.add_page_break()

        pdf_doc.close()

        if not extracted_any_text:
            return False, "لم يتم استخراج أي نص من الملف"

        doc.save(docx_path)
        return True, None

    except Exception as e:
        logging.error(f"OCR conversion error: {e}")
        return False, str(e)


def has_arabic_text_issues(text):
    """Check if extracted text has common Arabic encoding issues"""
    if not text:
        return True

    issue_patterns = [
        '\ufffd',
        '�',
        '\x00',
    ]

    for pattern in issue_patterns:
        if pattern in text:
            return True

    arabic_char_count = sum(
        1 for c in text
        if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
    total_chars = len(text.replace(' ', '').replace('\n', ''))

    if total_chars > 0 and arabic_char_count / total_chars < 0.1:
        has_arabic_looking = any('\u0600' <= c <= '\u06FF' for c in text)
        if not has_arabic_looking:
            return True

    return False


@app.route('/convert-pdf', methods=['POST'])
def convert_pdf():
    start_time = time.time()
    file_size = 0
    pdf_path = None
    docx_path = None

    try:
        if 'pdf' not in request.files:
            return jsonify({'error': 'لم يتم تحميل ملف PDF'}), 400

        pdf_file = request.files['pdf']
        mode = request.form.get('mode', 'original')
        extraction_method = request.form.get('extraction_method', 'auto')

        if pdf_file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        if not pdf_file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'الرجاء تحميل ملف PDF فقط'}), 400

        unique_id = str(uuid.uuid4())[:8]
        pdf_path = os.path.join(UPLOAD_FOLDER, f'input_{unique_id}.pdf')
        docx_path = os.path.join(UPLOAD_FOLDER, f'output_{unique_id}.docx')

        pdf_file.save(pdf_path)
        file_size = os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0

        conversion_success = False
        use_ocr = extraction_method == 'ocr' or extraction_method == 'auto'

        if extraction_method == 'auto' or extraction_method == 'standard':
            try:
                from pdf2docx import Converter
                cv = Converter(pdf_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()

                if os.path.exists(docx_path):
                    from docx import Document
                    test_doc = Document(docx_path)
                    all_text = ""
                    for para in test_doc.paragraphs:
                        all_text += para.text + " "

                    if extraction_method == 'auto' and has_arabic_text_issues(
                            all_text):
                        logging.info(
                            "Standard conversion has Arabic issues, trying OCR..."
                        )
                        safe_remove_file(docx_path)
                        use_ocr = True
                    else:
                        conversion_success = True
                        use_ocr = False

            except Exception as e:
                logging.warning(f"Standard PDF conversion failed: {e}")
                if extraction_method == 'standard':
                    safe_remove_file(pdf_path)
                    return jsonify(
                        {'error': f'فشل في تحويل PDF: {str(e)[:100]}'}), 500
                use_ocr = True

        if use_ocr and not conversion_success:
            logging.info("Using Tesseract OCR for Arabic text extraction (no Gemini API)...")
            success, error = convert_pdf_with_ocr(pdf_path, docx_path, use_tesseract=True)
            if success:
                conversion_success = True
            else:
                safe_remove_file(pdf_path)
                return jsonify({
                    'error':
                    f'فشل في استخراج النص: {error[:100] if error else "خطأ غير معروف"}'
                }), 500

        if mode == 'en':
            try:
                from docx import Document
                from deep_translator import GoogleTranslator

                doc = Document(docx_path)
                translator = GoogleTranslator(source='auto', target='en')

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        try:
                            translated_text = translator.translate(
                                paragraph.text)
                            if translated_text:
                                paragraph.text = translated_text
                        except Exception:
                            pass

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    try:
                                        translated_text = translator.translate(
                                            paragraph.text)
                                        if translated_text:
                                            paragraph.text = translated_text
                                    except Exception:
                                        pass

                doc.save(docx_path)

            except Exception as e:
                safe_remove_files(pdf_path, docx_path)
                return jsonify({'error':
                                f'فشل في الترجمة: {str(e)[:100]}'}), 500

        if not os.path.exists(docx_path):
            safe_remove_file(pdf_path)
            return jsonify({'error': 'فشل في إنشاء ملف Word'}), 500

        final_pdf = pdf_path
        final_docx = docx_path

        @after_this_request
        def cleanup(response):
            safe_remove_files(final_pdf, final_docx)
            return response

        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('pdf_converter', 'convert', 'success', 
                    duration_ms=duration_ms, file_size=file_size,
                    details=json.dumps({'mode': mode, 'method': extraction_method}))
        return send_file(
            docx_path,
            as_attachment=True,
            download_name='converted_document.docx',
            mimetype=
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        safe_remove_files(pdf_path, docx_path)
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('pdf_converter', 'convert', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(e))
        log_error('PDFConvertError', str(e), traceback.format_exc(), tool_name='pdf_converter')
        return jsonify({'error': f'خطأ غير متوقع: {str(e)[:100]}'}), 500


ALLOWED_VIDEO_EXTENSIONS = {
    'mp4', 'mov', 'avi', 'mkv', 'webm', 'flv', 'wmv', 'm4v'
}


def allowed_video(filename):
    return '.' in filename and filename.rsplit(
        '.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS


def extract_audio_from_video(video_path, output_audio_path):
    """Extract audio from video file using ffmpeg - optimized for speed"""
    try:
        ffmpeg_cmd = [
            'ffmpeg', '-y', '-i', video_path, '-vn', '-acodec', 'libmp3lame',
            '-q:a', '5', '-threads', '4', '-ar', '16000', '-ac', '1',
            output_audio_path
        ]
        subprocess.run(ffmpeg_cmd,
                       check=True,
                       capture_output=True,
                       timeout=7200)
        return output_audio_path
    except subprocess.CalledProcessError as e:
        logging.error(f"FFmpeg audio extraction error: {e}")
        return None
    except subprocess.TimeoutExpired:
        logging.error("FFmpeg audio extraction timed out")
        return None


@app.route('/transcribe-video', methods=['POST'])
def transcribe_video():
    start_time = time.time()
    file_size = 0
    temp_video = None
    temp_audio = None
    compressed_audio = None

    try:
        input_type = request.form.get('input_type', 'file')
        language = request.form.get('language', 'ar')
        if language not in LANGUAGE_NAMES:
            language = 'ar'

        logging.info(
            f"Video transcription started - input_type: {input_type}, language: {language}"
        )
        logging.info(
            f"Available AI providers: {get_available_keys_count()}/2"
        )

        if input_type == 'url':
            url = request.form.get('url', '').strip()

            if not url:
                return jsonify({'error': 'الرجاء إدخال رابط الفيديو'}), 400

            if not url.startswith(('http://', 'https://')):
                return jsonify({
                    'error':
                    'الرابط غير صالح. يجب أن يبدأ بـ http:// أو https://'
                }), 400

            unique_id = str(uuid.uuid4())[:8]
            temp_audio = os.path.join(UPLOAD_FOLDER,
                                      f'video_audio_{unique_id}.mp3')

            ydl_opts = {
                'format':
                'bestaudio/best',
                'outtmpl':
                temp_audio.replace('.mp3', '.%(ext)s'),
                'noplaylist':
                True,
                'quiet':
                True,
                'no_warnings':
                True,
                'socket_timeout':
                1800,
                'retries':
                15,
                'fragment_retries':
                15,
                'file_access_retries':
                10,
                'extractor_retries':
                10,
                'buffersize':
                1024,
                'http_chunk_size':
                10485760,
                'http_headers': {
                    'User-Agent':
                    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    'Accept':
                    'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                },
                'postprocessors': [{
                    'key': 'FFmpegExtractAudio',
                    'preferredcodec': 'mp3',
                    'preferredquality': '64',
                }],
            }

            if os.path.exists('cookies.txt'):
                ydl_opts['cookiefile'] = 'cookies.txt'

            max_download_attempts = 5
            download_success = False
            last_error = None

            for attempt in range(max_download_attempts):
                try:
                    logging.info(
                        f"Download attempt {attempt + 1}/{max_download_attempts} for URL: {url[:50]}..."
                    )
                    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                        ydl.download([url])
                    download_success = True
                    logging.info(
                        f"Download successful on attempt {attempt + 1}")
                    break
                except Exception as e:
                    last_error = e
                    error_msg = str(e).lower()
                    if 'private' in error_msg:
                        return jsonify(
                            {'error':
                             'هذا الفيديو خاص ولا يمكن الوصول إليه.'}), 400
                    elif 'unavailable' in error_msg or 'removed' in error_msg or 'deleted' in error_msg:
                        return jsonify(
                            {'error': 'الفيديو غير متاح أو تم حذفه.'}), 400
                    elif 'age' in error_msg or 'sign in' in error_msg:
                        return jsonify(
                            {'error': 'هذا الفيديو يتطلب تسجيل الدخول.'}), 400
                    else:
                        logging.warning(
                            f"Download attempt {attempt + 1} failed: {str(e)[:100]}"
                        )
                        if attempt < max_download_attempts - 1:
                            import time
                            wait_time = (attempt + 1) * 3
                            logging.info(
                                f"Waiting {wait_time}s before retry...")
                            time.sleep(wait_time)
                            continue

            if not download_success:
                return jsonify({
                    'error':
                    f'فشل في تحميل الفيديو بعد {max_download_attempts} محاولات. جرب مرة أخرى أو استخدم رابط آخر.'
                }), 400

            if not os.path.exists(temp_audio):
                for ext in ['mp3', 'wav', 'm4a', 'ogg', 'opus', 'webm']:
                    alt_file = temp_audio.replace('.mp3', f'.{ext}')
                    if os.path.exists(alt_file):
                        if ext != 'mp3':
                            new_temp_audio = os.path.join(
                                UPLOAD_FOLDER,
                                f'video_audio_{unique_id}_converted.mp3')
                            try:
                                ffmpeg_cmd = [
                                    'ffmpeg', '-y', '-i', alt_file, '-acodec',
                                    'libmp3lame', '-ab', '128k', new_temp_audio
                                ]
                                subprocess.run(ffmpeg_cmd,
                                               check=True,
                                               capture_output=True,
                                               timeout=7200)
                                safe_remove_file(alt_file)
                                temp_audio = new_temp_audio
                            except:
                                temp_audio = alt_file
                        else:
                            temp_audio = alt_file
                        break

            if not os.path.exists(temp_audio):
                return jsonify({'error':
                                'فشل في استخراج الصوت من الفيديو'}), 500

        else:
            if 'video' not in request.files:
                return jsonify({'error': 'الرجاء تحميل ملف فيديو'}), 400

            file = request.files['video']

            if file.filename == '':
                return jsonify({'error': 'لم يتم اختيار ملف'}), 400

            if not allowed_video(file.filename):
                return jsonify({
                    'error':
                    'نوع الملف غير مدعوم. الأنواع المدعومة: MP4, MOV, AVI, MKV, WEBM, FLV, WMV, M4V'
                }), 400

            unique_id = str(uuid.uuid4())[:8]
            ext = file.filename.rsplit('.', 1)[1].lower()
            temp_video = os.path.join(UPLOAD_FOLDER,
                                      f'video_{unique_id}.{ext}')
            temp_audio = os.path.join(UPLOAD_FOLDER,
                                      f'video_audio_{unique_id}.mp3')

            file.save(temp_video)

            extracted_audio = extract_audio_from_video(temp_video, temp_audio)

            if not extracted_audio or not os.path.exists(temp_audio):
                safe_remove_file(temp_video)
                return jsonify({
                    'error':
                    'فشل في استخراج الصوت من الفيديو. تأكد من أن الفيديو يحتوي على صوت.'
                }), 400

            safe_remove_file(temp_video)
            temp_video = None

        logging.info(
            f"Starting Gemini transcription for video audio: {temp_audio} with language: {language}"
        )

        try:
            text = transcribe_audio_with_gemini(temp_audio, language)
        except ValueError as ve:
            safe_remove_file(temp_audio)
            run_garbage_collection()
            return jsonify({'error': str(ve)}), 400
        except Exception as transcribe_error:
            logging.error(f"Transcription error: {transcribe_error}")
            safe_remove_files(temp_audio, temp_video)
            run_garbage_collection()
            return jsonify({
                'error':
                'فشل في تحويل الصوت إلى نص. حاول مرة أخرى أو استخدم ملف أقصر.'
            }), 500

        safe_remove_file(temp_audio)
        run_garbage_collection()

        duration_ms = int((time.time() - start_time) * 1000)
        if text:
            word_count = len(text.split())
            log_activity('video_transcription', 'transcribe', 'success', 
                        duration_ms=duration_ms, file_size=file_size,
                        details=json.dumps({'language': language, 'word_count': word_count, 'input_type': input_type}))
            return jsonify({
                'success': True,
                'text': text,
                'method': 'gemini_ai',
                'word_count': word_count,
                'keys_available': get_available_keys_count()
            })
        else:
            log_activity('video_transcription', 'transcribe', 'error', 
                        duration_ms=duration_ms, file_size=file_size, error_message='No text returned')
            return jsonify({
                'error':
                'فشل في تحويل الفيديو إلى نص. تأكد من إعداد مفاتيح Gemini API أو جرب ملف أقصر.'
            }), 500

    except ValueError as ve:
        safe_remove_files(temp_video, temp_audio)
        run_garbage_collection()
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('video_transcription', 'transcribe', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=str(ve))
        return jsonify({'error': str(ve)}), 400
    except Exception as e:
        logging.error(f"Video transcription error: {e}")
        safe_remove_files(temp_video, temp_audio)
        run_garbage_collection()
        duration_ms = int((time.time() - start_time) * 1000)
        error_message = str(e)
        log_activity('video_transcription', 'transcribe', 'error', 
                    duration_ms=duration_ms, file_size=file_size, error_message=error_message[:200])
        log_error('VideoTranscriptionError', error_message, traceback.format_exc(), tool_name='video_transcription')
        if 'timeout' in error_message.lower(
        ) or 'timed out' in error_message.lower():
            return jsonify({
                'error':
                'انتهت مهلة المعالجة. جرب ملف أقصر أو اتصال إنترنت أسرع.'
            }), 500
        elif 'memory' in error_message.lower():
            return jsonify({'error':
                            'الملف كبير جداً. جرب ملف أصغر حجماً.'}), 500
        elif 'quota' in error_message.lower() or 'rate' in error_message.lower(
        ):
            return jsonify(
                {'error': 'تم استنفاد حصة API. حاول مرة أخرى لاحقاً.'}), 500
        return jsonify(
            {'error': f'حدث خطأ أثناء المعالجة: {error_message[:80]}'}), 500


def apply_basic_format(doc: Document):
    """
    Basic formatting: margins 2.5 cm, font size 14, justified paragraphs, 1.5 line spacing.
    """
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(14)
    except Exception:
        pass

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in para.runs:
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"


def extract_formatting_from_ai(instructions=None, reference_image=None):
    """
    Use Gemini AI to understand formatting requirements from text or image.
    Returns a JSON-like dict with formatting specifications.
    Enhanced version with better prompt engineering for accurate formatting extraction.
    """
    import json
    
    default_format = {
        "font_name": "Times New Roman",
        "font_name_arabic": "Simplified Arabic",
        "font_size": 14,
        "line_spacing": 1.5,
        "alignment": "justify",
        "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
        "paragraph_spacing_before": 0,
        "paragraph_spacing_after": 6,
        "first_line_indent": 1.27,
        "heading_font_size": 16,
        "heading_bold": True,
        "heading_alignment": "center",
        "rtl_direction": True
    }
    
    result = None
    
    prompt = """أنت خبير متخصص في تنسيق المستندات الأكاديمية والرسمية. مهمتك تحليل متطلبات التنسيق وإرجاع مواصفات دقيقة بصيغة JSON.

قم بإرجاع JSON فقط بالضبط بهذا الشكل (بدون أي نص إضافي قبله أو بعده):
{
    "font_name": "Times New Roman",
    "font_name_arabic": "Simplified Arabic",
    "font_size": 14,
    "line_spacing": 1.5,
    "alignment": "justify",
    "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
    "paragraph_spacing_before": 0,
    "paragraph_spacing_after": 6,
    "first_line_indent": 1.27,
    "heading_font_size": 16,
    "heading_bold": true,
    "heading_alignment": "center",
    "rtl_direction": true
}

ملاحظات مهمة:
- font_name: الخط الإنجليزي (Times New Roman, Arial, Calibri)
- font_name_arabic: الخط العربي (Simplified Arabic, Traditional Arabic, Arial)
- font_size: حجم الخط بالنقاط (12, 14, 16)
- line_spacing: تباعد الأسطر (1.0, 1.15, 1.5, 2.0)
- alignment: محاذاة النص (justify, right, left, center)
- margins: الهوامش بالسنتيمتر
- first_line_indent: المسافة البادئة للسطر الأول بالسنتيمتر
- heading_alignment: محاذاة العناوين (center, right, left)
- rtl_direction: اتجاه النص من اليمين لليسار للعربية

"""
    
    if instructions:
        prompt += f"\n\nوصف التنسيق المطلوب:\n{instructions}\n\nقم بتحليل هذا الوصف وأرجع JSON فقط بالمواصفات المناسبة."
    else:
        prompt += "\n\nلم يتم تقديم تعليمات محددة. أرجع التنسيق الافتراضي للمستندات الأكاديمية العربية."
    
    try:
        if reference_image:
            prompt += "\n\nقم أيضاً بتحليل الصورة المرفقة واستخرج أسلوب التنسيق منها."
            result = call_gemini_vision(reference_image, prompt)
        else:
            result = call_gemini_text(prompt)
        
        if not result:
            logging.warning("AI returned no result, using default format")
            return default_format
        
        result_clean = result.strip()
        
        if "```json" in result_clean:
            start = result_clean.find("```json") + 7
            end = result_clean.find("```", start)
            if end > start:
                result_clean = result_clean[start:end].strip()
        elif "```" in result_clean:
            start = result_clean.find("```") + 3
            end = result_clean.find("```", start)
            if end > start:
                result_clean = result_clean[start:end].strip()
        
        if result_clean.startswith("{"):
            brace_count = 0
            json_end = 0
            for i, char in enumerate(result_clean):
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        json_end = i + 1
                        break
            if json_end > 0:
                result_clean = result_clean[:json_end]
        
        format_spec = json.loads(result_clean)
        
        for key in default_format:
            if key not in format_spec:
                format_spec[key] = default_format[key]
        
        logging.info(f"Successfully extracted format spec: {format_spec}")
        return format_spec
        
    except json.JSONDecodeError as e:
        logging.error(f"Failed to parse AI format response: {e}")
        logging.error(f"Raw response: {result[:500] if result else 'None'}")
        return default_format
    except Exception as e:
        logging.error(f"Error in extract_formatting_from_ai: {e}")
        return default_format


def apply_ai_format(doc: Document, format_spec: dict):
    """
    Apply AI-generated formatting specifications to a Word document.
    Enhanced version that properly applies all formatting to every paragraph and run.
    Supports documents up to 100+ pages.
    """
    from docx.shared import RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    alignment_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER
    }
    
    margins = format_spec.get("margins", {})
    for section in doc.sections:
        try:
            section.top_margin = Cm(float(margins.get("top", 2.5)))
            section.bottom_margin = Cm(float(margins.get("bottom", 2.5)))
            section.left_margin = Cm(float(margins.get("left", 2.5)))
            section.right_margin = Cm(float(margins.get("right", 2.5)))
        except Exception as e:
            logging.warning(f"Could not set margins: {e}")
    
    font_name = format_spec.get("font_name", "Times New Roman")
    font_name_arabic = format_spec.get("font_name_arabic", "Simplified Arabic")
    font_size = float(format_spec.get("font_size", 14))
    line_spacing = float(format_spec.get("line_spacing", 1.5))
    alignment_str = format_spec.get("alignment", "justify")
    alignment = alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.JUSTIFY)
    para_spacing_before = float(format_spec.get("paragraph_spacing_before", 0))
    para_spacing_after = float(format_spec.get("paragraph_spacing_after", 6))
    first_indent = float(format_spec.get("first_line_indent", 0))
    heading_size = float(format_spec.get("heading_font_size", 16))
    heading_bold = format_spec.get("heading_bold", True)
    heading_alignment_str = format_spec.get("heading_alignment", "center")
    heading_alignment = alignment_map.get(heading_alignment_str, WD_ALIGN_PARAGRAPH.CENTER)
    rtl_direction = format_spec.get("rtl_direction", True)
    
    try:
        for style_name in ["Normal", "Body Text", "Default Paragraph Font"]:
            try:
                style = doc.styles[style_name]
                style.font.name = font_name
                style.font.size = Pt(font_size)
                style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_arabic)
            except:
                pass
    except Exception as e:
        logging.warning(f"Could not set document styles: {e}")
    
    def is_arabic(text):
        """Check if text contains Arabic characters"""
        arabic_range = range(0x0600, 0x06FF + 1)
        return any(ord(char) in arabic_range for char in text)
    
    def is_heading(para):
        """Determine if a paragraph is a heading"""
        text = para.text.strip()
        if not text:
            return False
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            return True
        if len(text) < 80 and not text.endswith(('.', '،', '؟', '!', ':', ';')):
            if para.runs and any(run.bold for run in para.runs if run.text.strip()):
                return True
        if len(text) < 50 and text.isupper():
            return True
        return False
    
    total_paragraphs = len(doc.paragraphs)
    formatted_count = 0
    
    for idx, para in enumerate(doc.paragraphs):
        try:
            text = para.text.strip()
            if not text:
                continue
            
            formatted_count += 1
            is_head = is_heading(para)
            text_is_arabic = is_arabic(text)
            
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            pf.space_before = Pt(para_spacing_before)
            pf.space_after = Pt(para_spacing_after)
            
            if is_head:
                pf.alignment = heading_alignment
                pf.first_line_indent = Cm(0)
            else:
                pf.alignment = alignment
                if first_indent > 0:
                    pf.first_line_indent = Cm(first_indent)
            
            if rtl_direction and text_is_arabic:
                try:
                    pPr = para._element.get_or_add_pPr()
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    pPr.append(bidi)
                except:
                    pass
            
            current_font = font_name_arabic if text_is_arabic else font_name
            current_size = heading_size if is_head else font_size
            
            if para.runs:
                for run in para.runs:
                    try:
                        run.font.name = current_font
                        run.font.size = Pt(current_size)
                        
                        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                        run._element.rPr.rFonts.set(qn('w:cs'), font_name_arabic)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_arabic)
                        
                        if is_head:
                            run.font.bold = heading_bold
                    except Exception as run_error:
                        logging.debug(f"Could not format run: {run_error}")
            else:
                try:
                    run = para.add_run(text)
                    para.clear()
                    para.add_run(text)
                    for run in para.runs:
                        run.font.name = current_font
                        run.font.size = Pt(current_size)
                        if is_head:
                            run.font.bold = heading_bold
                except:
                    pass
                    
        except Exception as para_error:
            logging.warning(f"Error formatting paragraph {idx}: {para_error}")
            continue
    
    for table in doc.tables:
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue
                        text_is_arabic = is_arabic(text)
                        current_font = font_name_arabic if text_is_arabic else font_name
                        
                        for run in para.runs:
                            try:
                                run.font.name = current_font
                                run.font.size = Pt(font_size)
                            except:
                                pass
        except Exception as table_error:
            logging.warning(f"Error formatting table: {table_error}")
    
    logging.info(f"Successfully formatted {formatted_count}/{total_paragraphs} paragraphs")


def convert_pdf_to_docx_for_formatting(pdf_path, output_path):
    """Convert PDF to DOCX for formatting purposes."""
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except Exception as e:
        logging.error(f"PDF to DOCX conversion failed: {e}")
        return False


@app.route("/format-docx", methods=["POST"])
def format_docx():
    uploaded = request.files.get("file")

    if not uploaded or uploaded.filename == "":
        return jsonify({'error': 'من فضلك ارفع ملف DOCX'}), 400

    filename = uploaded.filename.lower()
    if not filename.endswith(".docx"):
        return jsonify({'error': 'النسخة الحالية تدعم ملفات Word بصيغة .docx فقط'}), 400

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "input.docx")
        out_path = os.path.join(tmpdir, "formatted.docx")

        uploaded.save(in_path)

        doc = Document(in_path)
        apply_basic_format(doc)
        doc.save(out_path)

        download_name = f"formatted_{uploaded.filename}"
        return send_file(out_path, as_attachment=True, download_name=download_name)


@app.route("/format-document-ai", methods=["POST"])
def format_document_ai():
    """
    AI-powered document formatting endpoint.
    Accepts: main document (DOCX/PDF), optional formatting instructions, optional reference file (image/DOCX/PDF)
    """
    start_time = time.time()
    file_size = 0
    
    try:
        uploaded = request.files.get("file")
        reference_file = request.files.get("reference_file")
        instructions = request.form.get("instructions", "").strip()
        
        if not uploaded or uploaded.filename == "":
            return jsonify({'error': 'من فضلك ارفع ملف للتنسيق'}), 400
        
        filename = uploaded.filename.lower()
        if not filename.endswith((".docx", ".pdf")):
            return jsonify({'error': 'يدعم ملفات Word (.docx) و PDF فقط'}), 400
        
        if not instructions and not reference_file:
            return jsonify({'error': 'يرجى إدخال وصف التنسيق أو رفع ملف مرجعي'}), 400
        
        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, "input" + os.path.splitext(uploaded.filename)[1])
            docx_path = os.path.join(tmpdir, "document.docx")
            out_path = os.path.join(tmpdir, "formatted.docx")
            
            uploaded.save(in_path)
            file_size = os.path.getsize(in_path) if os.path.exists(in_path) else 0
            
            if filename.endswith(".pdf"):
                logging.info("Converting PDF to DOCX for formatting...")
                if not convert_pdf_to_docx_for_formatting(in_path, docx_path):
                    duration_ms = int((time.time() - start_time) * 1000)
                    log_activity('document_formatter', 'format', 'error', duration_ms=duration_ms, 
                                file_size=file_size, error_message='PDF conversion failed')
                    return jsonify({'error': 'فشل في تحويل ملف PDF. جرب ملف Word مباشرة.'}), 500
            else:
                docx_path = in_path
            
            reference_image = None
            if reference_file and reference_file.filename:
                ref_filename = reference_file.filename.lower()
                ref_path = os.path.join(tmpdir, "reference" + os.path.splitext(reference_file.filename)[1])
                reference_file.save(ref_path)
                
                if ref_filename.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
                    reference_image = ref_path
                elif ref_filename.endswith(".pdf"):
                    try:
                        import fitz
                        pdf_doc = fitz.open(ref_path)
                        page = pdf_doc[0]
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img_path = os.path.join(tmpdir, "ref_screenshot.png")
                        pix.save(img_path)
                        pdf_doc.close()
                        reference_image = img_path
                        logging.info("Successfully extracted image from reference PDF")
                    except Exception as e:
                        logging.error(f"Failed to extract image from reference PDF: {e}")
                        instructions += "\n\nالملف المرجعي هو PDF. استخدم تنسيق بحثي أكاديمي مشابه."
                elif ref_filename.endswith(".docx"):
                    try:
                        ref_doc = Document(ref_path)
                        ref_format_info = []
                        
                        if ref_doc.sections:
                            section = ref_doc.sections[0]
                            margins = {
                                "top": round(section.top_margin.cm, 1) if section.top_margin else 2.5,
                                "bottom": round(section.bottom_margin.cm, 1) if section.bottom_margin else 2.5,
                                "left": round(section.left_margin.cm, 1) if section.left_margin else 2.5,
                                "right": round(section.right_margin.cm, 1) if section.right_margin else 2.5
                            }
                            ref_format_info.append(f"الهوامش: أعلى {margins['top']}سم، أسفل {margins['bottom']}سم، يسار {margins['left']}سم، يمين {margins['right']}سم")
                        
                        for para in ref_doc.paragraphs[:5]:
                            if para.runs:
                                run = para.runs[0]
                                if run.font.name:
                                    ref_format_info.append(f"الخط: {run.font.name}")
                                if run.font.size:
                                    ref_format_info.append(f"حجم الخط: {run.font.size.pt} نقطة")
                                break
                        
                        for para in ref_doc.paragraphs[:5]:
                            if para.paragraph_format.line_spacing:
                                ref_format_info.append(f"تباعد الأسطر: {para.paragraph_format.line_spacing}")
                                break
                        
                        if ref_format_info:
                            instructions += "\n\nالتنسيق المستخرج من الملف المرجعي:\n" + "\n".join(ref_format_info)
                        else:
                            instructions += "\n\nالملف المرجعي هو مستند Word. استخدم تنسيق بحثي أكاديمي مشابه."
                        
                        logging.info(f"Extracted formatting from reference DOCX: {ref_format_info}")
                    except Exception as e:
                        logging.error(f"Failed to extract formatting from reference DOCX: {e}")
                        instructions += "\n\nالملف المرجعي هو مستند Word. استخدم تنسيق بحثي أكاديمي مشابه."
            
            logging.info(f"Extracting formatting specs from AI... Instructions: {instructions[:100] if instructions else 'None'}, Reference: {bool(reference_image)}")
            format_spec = extract_formatting_from_ai(instructions, reference_image)
            logging.info(f"AI format spec: {format_spec}")
            
            doc = Document(docx_path)
            apply_ai_format(doc, format_spec)
            doc.save(out_path)
            
            original_name = os.path.splitext(uploaded.filename)[0]
            download_name = f"formatted_{original_name}.docx"
            
            output_size = os.path.getsize(out_path) if os.path.exists(out_path) else 0
            duration_ms = int((time.time() - start_time) * 1000)
            log_activity('document_formatter', 'format', 'success', duration_ms=duration_ms, 
                        file_size=file_size, details=json.dumps({'output_size': output_size}))
            
            return send_file(out_path, as_attachment=True, download_name=download_name)
    
    except Exception as e:
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('document_formatter', 'format', 'error', duration_ms=duration_ms, 
                    file_size=file_size, error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='document_formatter')
        logging.error(f"Document formatting error: {e}")
        return jsonify({'error': f'خطأ في التنسيق: {str(e)}'}), 500


# Extract audio from YouTube for background
@app.route('/get-audio', methods=['POST'])
def get_audio():
    start_time = time.time()
    file_size = 0
    
    data = request.get_json()
    url = data.get('url')
    
    if not url:
        return jsonify({'error': 'No URL provided'}), 400
    
    try:
        temp_dir = tempfile.gettempdir()
        ydl_opts = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '128',
            }],
            'outtmpl': os.path.join(temp_dir, 'arkan_audio_%(id)s'),
            'quiet': True,
            'no_warnings': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            audio_file = os.path.join(temp_dir, f"arkan_audio_{info['id']}.mp3")
            
            if os.path.exists(audio_file):
                file_size = os.path.getsize(audio_file)
                duration_ms = int((time.time() - start_time) * 1000)
                log_activity('audio_download', 'extract', 'success', duration_ms=duration_ms, 
                            file_size=file_size, details=json.dumps({'video_id': info['id']}))
                
                @after_this_request
                def cleanup(response):
                    safe_remove_file(audio_file)
                    return response
                
                return send_file(audio_file, mimetype='audio/mpeg', as_attachment=False, download_name=f"audio_{info['id']}.mp3")
            else:
                duration_ms = int((time.time() - start_time) * 1000)
                log_activity('audio_download', 'extract', 'error', duration_ms=duration_ms, 
                            error_message='Failed to extract audio')
                return jsonify({'error': 'Failed to extract audio'}), 500
    
    except Exception as e:
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('audio_download', 'extract', 'error', duration_ms=duration_ms, 
                    file_size=file_size, error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='audio_download')
        logging.error(f"Audio extraction error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# =============================================================================
# ADMIN DASHBOARD ROUTES
# =============================================================================
# Protected admin dashboard for Gemini API key management
# Access only via: /admin/keys (requires admin login)
# =============================================================================

ADMIN_DASHBOARD_PATH = '/admin/keys'


@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    """Admin login page."""
    if current_user.is_authenticated:
        return redirect(url_for('admin_unified_dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = AdminUser.query.filter_by(username=username).first()
        if user and user.check_password(password):
            user.last_login = datetime.utcnow()
            db.session.commit()
            login_user(user)
            logging.info(f"[Admin] User '{username}' logged in successfully")
            return redirect(url_for('admin_unified_dashboard'))
        
        flash('اسم المستخدم أو كلمة المرور غير صحيحة', 'error')
    
    return render_template('admin_login.html')


@app.route('/admin/logout')
@login_required
def admin_logout():
    """Admin logout."""
    logout_user()
    return redirect(url_for('admin_login'))


@app.route(ADMIN_DASHBOARD_PATH)
@login_required
def admin_dashboard():
    """Redirect old admin/keys route to new unified admin dashboard."""
    return redirect(url_for('admin_unified_dashboard'))


@app.route('/admin/api/clear-cache', methods=['POST'])
@login_required
def admin_clear_cache():
    """API endpoint to clear the AI response cache."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    ai_manager.cache.clear()
    
    return jsonify({
        'success': True,
        'message': 'تم مسح ذاكرة التخزين المؤقت بنجاح'
    })


@app.route('/admin/api/stats')
@login_required
def admin_api_stats():
    """API endpoint to get current stats (for dashboard refresh)."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    stats = ai_manager.get_stats()
    
    providers_info = [
        {
            'name': 'Groq',
            'type': 'LLM + Audio',
            'status': 'active' if ai_manager.groq.is_configured else 'not_configured',
            'is_configured': ai_manager.groq.is_configured
        },
        {
            'name': 'HuggingFace', 
            'type': 'Vision',
            'status': 'active' if ai_manager.huggingface.is_configured else 'not_configured',
            'is_configured': ai_manager.huggingface.is_configured
        }
    ]
    
    return jsonify({
        'stats': stats,
        'providers_info': providers_info,
        'cache_stats': ai_manager.cache.stats()
    })


@app.route('/admin/api/active-sessions')
@login_required
def admin_api_active_sessions():
    """API endpoint to get active sessions for dashboard."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    five_minutes_ago = datetime.utcnow() - timedelta(minutes=5)
    one_hour_ago = datetime.utcnow() - timedelta(hours=1)
    
    ActiveSession.query.filter(ActiveSession.last_seen < five_minutes_ago).update({'is_active': False})
    db.session.commit()
    
    active_now = ActiveSession.query.filter(ActiveSession.last_seen >= five_minutes_ago).count()
    active_hour = ActiveSession.query.filter(ActiveSession.last_seen >= one_hour_ago).count()
    total_sessions = ActiveSession.query.count()
    
    recent_sessions = ActiveSession.query.filter(
        ActiveSession.last_seen >= one_hour_ago
    ).order_by(ActiveSession.last_seen.desc()).limit(50).all()
    
    device_stats = db.session.query(
        ActiveSession.device_type,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.device_type).all()
    
    browser_stats = db.session.query(
        ActiveSession.browser,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.browser).all()
    
    sessions_data = []
    for s in recent_sessions:
        sessions_data.append({
            'session_id': s.session_id[:8] + '...',
            'ip_address': s.ip_address,
            'device_type': s.device_type,
            'browser': s.browser,
            'os_name': s.os_name,
            'page_views': s.page_views,
            'first_seen': s.first_seen.isoformat() if s.first_seen else None,
            'last_seen': s.last_seen.isoformat() if s.last_seen else None,
            'is_active': s.is_active
        })
    
    return jsonify({
        'active_now': active_now,
        'active_hour': active_hour,
        'total_sessions': total_sessions,
        'sessions': sessions_data,
        'device_stats': dict(device_stats),
        'browser_stats': dict(browser_stats)
    })


@app.route('/admin')
@login_required
def admin_unified_dashboard():
    """Unified admin dashboard with all features."""
    if not current_user.is_admin:
        flash('ليس لديك صلاحية الوصول إلى لوحة التحكم', 'error')
        return redirect(url_for('admin_login'))
    
    stats = ai_manager.get_stats()
    today = date.today()
    one_hour_ago = datetime.utcnow() - timedelta(hours=1)
    five_minutes_ago = datetime.utcnow() - timedelta(minutes=5)
    
    active_now = ActiveSession.query.filter(ActiveSession.last_seen >= five_minutes_ago).count()
    active_hour = ActiveSession.query.filter(ActiveSession.last_seen >= one_hour_ago).count()
    
    hourly_data = []
    groq_hourly_data = []
    for hour in range(24):
        hourly_stats = HourlyStats.query.filter_by(date=today, hour=hour).first()
        if hourly_stats:
            hourly_data.append({
                'hour': hour,
                'total': hourly_stats.total_requests,
                'success': hourly_stats.successful_requests,
                'failed': hourly_stats.failed_requests
            })
            groq_hourly_data.append({
                'hour': hour,
                'groq': hourly_stats.groq_requests or 0,
                'huggingface': hourly_stats.huggingface_requests or 0
            })
        else:
            in_memory_hourly = stats.get('hourly', {}).get(hour, {})
            hourly_data.append({
                'hour': hour,
                'total': in_memory_hourly.get('total', 0),
                'success': in_memory_hourly.get('success', 0),
                'failed': in_memory_hourly.get('failed', 0)
            })
            groq_hourly_data.append({
                'hour': hour,
                'groq': 0,
                'huggingface': 0
            })
    
    groq_usage_logs = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.utcnow() - timedelta(hours=24),
        AIUsageLog.provider == 'groq'
    ).all()
    for log in groq_usage_logs:
        if log.timestamp:
            hour = log.timestamp.hour
            groq_hourly_data[hour]['groq'] += 1
    
    hf_usage_logs = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.utcnow() - timedelta(hours=24),
        AIUsageLog.provider == 'huggingface'
    ).all()
    for log in hf_usage_logs:
        if log.timestamp:
            hour = log.timestamp.hour
            groq_hourly_data[hour]['huggingface'] += 1
    
    groq_errors_today = ErrorLog.query.filter(
        ErrorLog.timestamp >= datetime.combine(today, datetime.min.time()),
        ErrorLog.provider == 'groq'
    ).count()
    hf_errors_today = ErrorLog.query.filter(
        ErrorLog.timestamp >= datetime.combine(today, datetime.min.time()),
        ErrorLog.provider == 'huggingface'
    ).count()
    
    provider_errors = {
        'groq': groq_errors_today,
        'huggingface': hf_errors_today
    }
    
    groq_success_today = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.combine(today, datetime.min.time()),
        AIUsageLog.provider == 'groq',
        AIUsageLog.success == True
    ).count()
    groq_failed_today = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.combine(today, datetime.min.time()),
        AIUsageLog.provider == 'groq',
        AIUsageLog.success == False
    ).count()
    hf_success_today = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.combine(today, datetime.min.time()),
        AIUsageLog.provider == 'huggingface',
        AIUsageLog.success == True
    ).count()
    hf_failed_today = AIUsageLog.query.filter(
        AIUsageLog.timestamp >= datetime.combine(today, datetime.min.time()),
        AIUsageLog.provider == 'huggingface',
        AIUsageLog.success == False
    ).count()
    
    provider_stats = {
        'groq': {'success': groq_success_today, 'failed': groq_failed_today, 'errors': groq_errors_today},
        'huggingface': {'success': hf_success_today, 'failed': hf_failed_today, 'errors': hf_errors_today}
    }
    
    request_type_data = {
        'LLM (نص)': stats.get('llm_requests', 0),
        'Whisper (صوت)': stats.get('whisper_requests', 0),
        'Vision (صور)': stats.get('vision_requests', 0),
        'مخزنة مؤقتاً': stats.get('cached_requests', 0)
    }
    
    device_stats = db.session.query(
        ActiveSession.device_type,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.device_type).all()
    
    browser_stats = db.session.query(
        ActiveSession.browser,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.browser).all()
    
    os_stats = db.session.query(
        ActiveSession.os_name,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.os_name).all()
    
    device_data = dict(device_stats) if device_stats else {'desktop': 0, 'mobile': 0, 'tablet': 0}
    browser_data = dict(browser_stats) if browser_stats else {'Chrome': 0, 'Firefox': 0, 'Safari': 0}
    os_data = dict(os_stats) if os_stats else {'Windows': 0, 'macOS': 0, 'Linux': 0}
    
    tool_stats_today = ToolStats.query.filter_by(date=today).order_by(ToolStats.usage_count.desc()).limit(10).all()
    top_tools = [{'name': t.tool_name, 'count': t.usage_count} for t in tool_stats_today]
    if not top_tools:
        top_tools = [
            {'name': 'تحويل الصوت للنص', 'count': 0},
            {'name': 'كاشف الأنمي', 'count': 0},
            {'name': 'كاشف البودكاست', 'count': 0}
        ]
    
    recent_activities = ActivityLog.query.order_by(ActivityLog.timestamp.desc()).limit(20).all()
    activities = [{
        'id': a.id,
        'timestamp': a.timestamp.isoformat() if a.timestamp else None,
        'tool_name': a.tool_name,
        'action': a.action,
        'status': a.status,
        'duration_ms': a.duration_ms,
        'session_id': a.session_id[:8] + '...' if a.session_id else 'N/A'
    } for a in recent_activities]
    
    recent_errors = ErrorLog.query.filter_by(resolved=False).order_by(ErrorLog.timestamp.desc()).limit(10).all()
    errors = [{
        'id': e.id,
        'timestamp': e.timestamp.isoformat() if e.timestamp else None,
        'error_type': e.error_type,
        'error_message': e.error_message[:100] if e.error_message else '',
        'tool_name': e.tool_name,
        'resolved': e.resolved
    } for e in recent_errors]
    
    providers_info = [
        {
            'name': 'Groq',
            'type': 'LLM + Audio',
            'status': 'active' if ai_manager.groq.is_configured else 'not_configured',
            'is_configured': ai_manager.groq.is_configured,
            'models': ['Llama 3.3 70B', 'Whisper Large V3'],
            'avg_latency': stats.get('groq_latency_ms', 0),
            'stats': {
                'total': provider_stats['groq']['success'] + provider_stats['groq']['failed'],
                'success': provider_stats['groq']['success'],
                'failed': provider_stats['groq']['failed'],
                'errors': provider_stats['groq']['errors']
            }
        },
        {
            'name': 'HuggingFace',
            'type': 'Vision',
            'status': 'active' if ai_manager.huggingface.is_configured else 'not_configured',
            'is_configured': ai_manager.huggingface.is_configured,
            'models': ['BLIP Image Captioning', 'ViLT VQA'],
            'avg_latency': stats.get('huggingface_latency_ms', 0),
            'stats': {
                'total': provider_stats['huggingface']['success'] + provider_stats['huggingface']['failed'],
                'success': provider_stats['huggingface']['success'],
                'failed': provider_stats['huggingface']['failed'],
                'errors': provider_stats['huggingface']['errors']
            }
        }
    ]
    
    latency = {
        'groq_llm': stats.get('groq_llm_latency_ms', 0),
        'groq_whisper': stats.get('groq_whisper_latency_ms', 0),
        'hf_vision': stats.get('huggingface_latency_ms', 0)
    }
    
    total_sessions = ActiveSession.query.count()
    page_views_today = db.session.query(db.func.sum(ActiveSession.page_views)).filter(
        ActiveSession.last_seen >= datetime.utcnow() - timedelta(days=1)
    ).scalar() or 0
    
    recent_sessions = ActiveSession.query.filter(
        ActiveSession.last_seen >= one_hour_ago
    ).order_by(ActiveSession.last_seen.desc()).limit(50).all()
    sessions_list = [{
        'session_id': s.session_id[:8] + '...',
        'ip_address': s.ip_address,
        'device_type': s.device_type,
        'browser': s.browser,
        'os_name': s.os_name,
        'page_views': s.page_views,
        'last_seen': s.last_seen.isoformat() if s.last_seen else None,
        'is_active': s.last_seen >= five_minutes_ago if s.last_seen else False
    } for s in recent_sessions]
    
    all_admins = AdminUser.query.all()
    admins_list = [{
        'id': a.id,
        'username': a.username,
        'last_login': a.last_login.isoformat() if a.last_login else None,
        'is_current': a.id == current_user.id
    } for a in all_admins]
    
    tool_names = db.session.query(ActivityLog.tool_name).distinct().all()
    tool_names = [t[0] for t in tool_names if t[0]]
    
    tool_stats_all = ToolStats.query.filter_by(date=today).order_by(ToolStats.usage_count.desc()).all()
    tool_stats_data = [{
        'name': t.tool_name,
        'usage_count': t.usage_count,
        'success_count': t.success_count,
        'error_count': t.error_count,
        'avg_duration_ms': t.avg_duration_ms
    } for t in tool_stats_all]
    
    settings = {
        'rate_limit': 15,
        'max_audio_minutes': MAX_AUDIO_DURATION_MINUTES,
        'cache_ttl': 3600
    }
    
    return render_template('admin.html',
        stats={
            'total_today': stats.get('total_today', 0),
            'success_today': stats.get('success_today', 0),
            'failed_today': stats.get('failed_today', 0),
            'cached_today': stats.get('cached_requests', 0)
        },
        active_sessions={'now': active_now, 'hour': active_hour, 'total': total_sessions},
        hourly_data=hourly_data,
        groq_hourly_data=groq_hourly_data,
        provider_errors=provider_errors,
        provider_stats=provider_stats,
        request_type_data=request_type_data,
        device_data=device_data,
        browser_data=browser_data,
        os_data=os_data,
        top_tools=top_tools,
        activities=recent_activities,
        errors=recent_errors,
        providers_info=providers_info,
        sessions=recent_sessions,
        admins=all_admins,
        cache_stats=ai_manager.cache.stats(),
        latency=latency,
        page_views_today=page_views_today,
        tool_names=tool_names,
        tool_stats=tool_stats_data,
        settings=settings
    )


@app.route('/admin/api/dashboard-stats')
@login_required
def admin_api_dashboard_stats():
    """API endpoint for dashboard stats refresh."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    stats = ai_manager.get_stats()
    five_minutes_ago = datetime.utcnow() - timedelta(minutes=5)
    one_hour_ago = datetime.utcnow() - timedelta(hours=1)
    
    active_now = ActiveSession.query.filter(ActiveSession.last_seen >= five_minutes_ago).count()
    active_hour = ActiveSession.query.filter(ActiveSession.last_seen >= one_hour_ago).count()
    
    return jsonify({
        'stats': {
            'total_today': stats.get('total_today', 0),
            'success_today': stats.get('success_today', 0),
            'failed_today': stats.get('failed_today', 0),
            'cached_today': stats.get('cached_requests', 0)
        },
        'active_sessions': {'now': active_now, 'hour': active_hour},
        'cache_stats': ai_manager.cache.stats()
    })


@app.route('/admin/api/sessions')
@login_required
def admin_api_sessions():
    """API endpoint for sessions tab data."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    five_minutes_ago = datetime.utcnow() - timedelta(minutes=5)
    one_hour_ago = datetime.utcnow() - timedelta(hours=1)
    
    active_now = ActiveSession.query.filter(ActiveSession.last_seen >= five_minutes_ago).count()
    active_hour = ActiveSession.query.filter(ActiveSession.last_seen >= one_hour_ago).count()
    total_sessions = ActiveSession.query.count()
    
    recent_sessions = ActiveSession.query.filter(
        ActiveSession.last_seen >= one_hour_ago
    ).order_by(ActiveSession.last_seen.desc()).limit(50).all()
    
    sessions_data = [{
        'session_id': s.session_id[:8] + '...',
        'ip_address': s.ip_address,
        'device_type': s.device_type,
        'browser': s.browser,
        'os_name': s.os_name,
        'page_views': s.page_views,
        'first_seen': s.first_seen.isoformat() if s.first_seen else None,
        'last_seen': s.last_seen.isoformat() if s.last_seen else None,
        'is_active': s.last_seen >= five_minutes_ago if s.last_seen else False
    } for s in recent_sessions]
    
    device_stats = db.session.query(
        ActiveSession.device_type,
        db.func.count(ActiveSession.id)
    ).filter(ActiveSession.last_seen >= one_hour_ago).group_by(ActiveSession.device_type).all()
    
    return jsonify({
        'active_now': active_now,
        'active_hour': active_hour,
        'total_sessions': total_sessions,
        'sessions': sessions_data,
        'device_stats': dict(device_stats)
    })


@app.route('/admin/api/activities')
@login_required
def admin_api_activities():
    """API endpoint for activity logs with filters."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    tool_filter = request.args.get('tool', '')
    status_filter = request.args.get('status', '')
    limit = min(int(request.args.get('limit', 50)), 200)
    
    query = ActivityLog.query
    if tool_filter:
        query = query.filter(ActivityLog.tool_name == tool_filter)
    if status_filter:
        query = query.filter(ActivityLog.status == status_filter)
    
    activities = query.order_by(ActivityLog.timestamp.desc()).limit(limit).all()
    
    return jsonify({
        'activities': [{
            'id': a.id,
            'timestamp': a.timestamp.isoformat() if a.timestamp else None,
            'tool_name': a.tool_name,
            'action': a.action,
            'status': a.status,
            'duration_ms': a.duration_ms,
            'session_id': a.session_id[:8] + '...' if a.session_id else 'N/A',
            'details': a.details
        } for a in activities],
        'total': ActivityLog.query.count()
    })


@app.route('/admin/api/errors')
@login_required
def admin_api_errors():
    """API endpoint for error logs."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    show_resolved = request.args.get('resolved', 'false') == 'true'
    limit = min(int(request.args.get('limit', 50)), 200)
    
    query = ErrorLog.query
    if not show_resolved:
        query = query.filter_by(resolved=False)
    
    errors = query.order_by(ErrorLog.timestamp.desc()).limit(limit).all()
    
    return jsonify({
        'errors': [{
            'id': e.id,
            'timestamp': e.timestamp.isoformat() if e.timestamp else None,
            'error_type': e.error_type,
            'error_message': e.error_message,
            'stack_trace': e.stack_trace[:500] if e.stack_trace else None,
            'tool_name': e.tool_name,
            'provider': e.provider,
            'resolved': e.resolved
        } for e in errors],
        'total': ErrorLog.query.filter_by(resolved=False).count()
    })


@app.route('/admin/api/errors/<int:error_id>/resolve', methods=['POST'])
@login_required
def admin_api_resolve_error(error_id):
    """API endpoint to mark an error as resolved."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    error = ErrorLog.query.get(error_id)
    if not error:
        return jsonify({'error': 'Error not found'}), 404
    
    error.resolved = True
    error.resolved_at = datetime.utcnow()
    error.resolution_notes = request.json.get('notes', '')
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/admin/api/settings', methods=['GET', 'POST'])
@login_required
def admin_api_settings():
    """API endpoint for AI settings management."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    if request.method == 'GET':
        return jsonify({
            'rate_limit': 15,
            'max_audio_minutes': MAX_AUDIO_DURATION_MINUTES,
            'cache_ttl': 3600
        })
    
    data = request.json
    logging.info(f"[Admin] Settings update: {data}")
    
    return jsonify({'success': True, 'message': 'تم حفظ الإعدادات'})


@app.route('/admin/api/admins', methods=['POST'])
@login_required
def admin_api_add_admin():
    """API endpoint to add new admin user."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '')
    
    if not username or not password:
        return jsonify({'error': 'اسم المستخدم وكلمة المرور مطلوبان'}), 400
    
    if len(password) < 6:
        return jsonify({'error': 'كلمة المرور يجب أن تكون 6 أحرف على الأقل'}), 400
    
    existing = AdminUser.query.filter_by(username=username).first()
    if existing:
        return jsonify({'error': 'اسم المستخدم موجود مسبقاً'}), 400
    
    new_admin = AdminUser(username=username, is_admin=True)
    new_admin.set_password(password)
    db.session.add(new_admin)
    db.session.commit()
    
    logging.info(f"[Admin] New admin user created: {username}")
    
    return jsonify({'success': True, 'message': 'تم إضافة المسؤول بنجاح'})


@app.route('/admin/api/admins/<int:admin_id>', methods=['DELETE'])
@login_required
def admin_api_delete_admin(admin_id):
    """API endpoint to delete an admin user."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    if admin_id == current_user.id:
        return jsonify({'error': 'لا يمكنك حذف حسابك الخاص'}), 400
    
    admin = AdminUser.query.get(admin_id)
    if not admin:
        return jsonify({'error': 'المسؤول غير موجود'}), 404
    
    admin_count = AdminUser.query.filter_by(is_admin=True).count()
    if admin_count <= 1:
        return jsonify({'error': 'يجب أن يبقى مسؤول واحد على الأقل'}), 400
    
    username = admin.username
    db.session.delete(admin)
    db.session.commit()
    
    logging.info(f"[Admin] Admin user deleted: {username}")
    
    return jsonify({'success': True, 'message': 'تم حذف المسؤول'})


@app.route('/admin/api/change-password', methods=['POST'])
@login_required
def admin_api_change_password():
    """API endpoint to change current admin's password."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json
    current_password = data.get('current_password', '')
    new_password = data.get('new_password', '')
    
    if not current_password or not new_password:
        return jsonify({'error': 'جميع الحقول مطلوبة'}), 400
    
    if not current_user.check_password(current_password):
        return jsonify({'error': 'كلمة المرور الحالية غير صحيحة'}), 400
    
    if len(new_password) < 6:
        return jsonify({'error': 'كلمة المرور الجديدة يجب أن تكون 6 أحرف على الأقل'}), 400
    
    current_user.set_password(new_password)
    db.session.commit()
    
    logging.info(f"[Admin] Password changed for user: {current_user.username}")
    
    return jsonify({'success': True, 'message': 'تم تغيير كلمة المرور بنجاح'})


# =============================================================================
# VIDEO TO MP3 CONVERTER TOOL
# =============================================================================
# Converts video files to MP3 audio using ffmpeg only (no AI)
# Supports: mp4, mkv, mov, avi, webm, flv, wmv
# =============================================================================

ALLOWED_VIDEO_EXTENSIONS_CONVERT = {'mp4', 'mkv', 'mov', 'avi', 'webm', 'flv', 'wmv', 'm4v', '3gp'}
MAX_VIDEO_SIZE_MB_CONVERT = 100
MAX_VIDEO_DURATION_CONVERT = 30  # minutes

def allowed_video_convert(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS_CONVERT

def get_video_duration_seconds(filepath):
    """Get video duration in seconds using ffprobe."""
    try:
        cmd = [
            'ffprobe', '-v', 'error', '-show_entries', 'format=duration',
            '-of', 'default=noprint_wrappers=1:nokey=1', filepath
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return float(result.stdout.strip())
    except Exception as e:
        logging.warning(f"Could not get video duration: {e}")
    return 0


@app.route('/api/video-to-mp3', methods=['POST'])
def video_to_mp3():
    """Convert uploaded video to MP3 audio."""
    start_time = time.time()
    temp_video = None
    temp_audio = None
    
    try:
        if 'video' not in request.files:
            return jsonify({'error': 'لم يتم رفع أي ملف فيديو'}), 400
        
        video_file = request.files['video']
        if video_file.filename == '':
            return jsonify({'error': 'اسم الملف فارغ'}), 400
        
        if not allowed_video_convert(video_file.filename):
            return jsonify({'error': 'صيغة الفيديو غير مدعومة. الصيغ المدعومة: MP4, MKV, MOV, AVI, WEBM, FLV, WMV'}), 400
        
        # Check file size
        video_file.seek(0, 2)
        file_size = video_file.tell()
        video_file.seek(0)
        
        if file_size > MAX_VIDEO_SIZE_MB_CONVERT * 1024 * 1024:
            return jsonify({'error': f'حجم الملف أكبر من الحد الأقصى المسموح ({MAX_VIDEO_SIZE_MB_CONVERT}MB)'}), 400
        
        # Get quality setting
        quality = request.form.get('quality', 'normal')
        bitrate = '192k' if quality == 'high' else '128k'
        
        # Save video temporarily
        ext = video_file.filename.rsplit('.', 1)[1].lower()
        temp_video = os.path.join(tempfile.gettempdir(), f"video_convert_{uuid.uuid4().hex}.{ext}")
        video_file.save(temp_video)
        
        # Check duration
        duration = get_video_duration_seconds(temp_video)
        if duration > MAX_VIDEO_DURATION_CONVERT * 60:
            safe_remove_file(temp_video)
            return jsonify({'error': f'مدة الفيديو أطول من الحد الأقصى المسموح ({MAX_VIDEO_DURATION_CONVERT} دقيقة)'}), 400
        
        # Convert to MP3
        temp_audio = os.path.join(tempfile.gettempdir(), f"audio_convert_{uuid.uuid4().hex}.mp3")
        
        cmd = [
            'ffmpeg', '-i', temp_video,
            '-vn', '-acodec', 'libmp3lame',
            '-ab', bitrate,
            '-ar', '44100',
            '-y', temp_audio
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        
        if result.returncode != 0:
            logging.error(f"FFmpeg error: {result.stderr}")
            safe_remove_file(temp_video)
            return jsonify({'error': 'فشل في تحويل الفيديو إلى صوت'}), 500
        
        safe_remove_file(temp_video)
        
        if not os.path.exists(temp_audio):
            return jsonify({'error': 'فشل في إنشاء ملف الصوت'}), 500
        
        duration_ms = int((time.time() - start_time) * 1000)
        log_activity('video_to_mp3', 'convert', 'success', duration_ms=duration_ms, file_size=file_size)
        
        original_name = os.path.splitext(video_file.filename)[0]
        
        @after_this_request
        def cleanup(response):
            safe_remove_file(temp_audio)
            return response
        
        return send_file(
            temp_audio,
            mimetype='audio/mpeg',
            as_attachment=True,
            download_name=f"{original_name}.mp3"
        )
        
    except subprocess.TimeoutExpired:
        safe_remove_files(temp_video, temp_audio)
        log_activity('video_to_mp3', 'convert', 'error', error_message='Timeout')
        return jsonify({'error': 'انتهى الوقت المسموح للتحويل. جرب ملف أصغر.'}), 408
        
    except Exception as e:
        safe_remove_files(temp_video, temp_audio)
        log_activity('video_to_mp3', 'convert', 'error', error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='video_to_mp3')
        logging.error(f"Video to MP3 error: {e}")
        return jsonify({'error': f'حدث خطأ: {str(e)}'}), 500


# =============================================================================
# MUSIC REMOVAL FROM AUDIO/VIDEO TOOL
# =============================================================================
# Removes background music to isolate human voice using ffmpeg filters
# Uses frequency filtering (no AI/ML required)
# Supports both audio and video files
# =============================================================================

ALLOWED_AUDIO_EXTENSIONS_CLEAN = {'mp3', 'wav', 'm4a', 'ogg', 'flac', 'aac', 'wma', 'webm'}
ALLOWED_VIDEO_EXTENSIONS_CLEAN = {'mp4', 'mkv', 'mov', 'avi', 'webm', 'flv', 'wmv'}
ALLOWED_MEDIA_EXTENSIONS_CLEAN = ALLOWED_AUDIO_EXTENSIONS_CLEAN | ALLOWED_VIDEO_EXTENSIONS_CLEAN
MAX_MEDIA_SIZE_MB_CLEAN = 100
MAX_MEDIA_DURATION_CLEAN = 30  # minutes

def allowed_media_clean(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_MEDIA_EXTENSIONS_CLEAN

def is_video_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS_CLEAN

def get_music_removal_filter(mode):
    """Get the ffmpeg audio filter based on processing mode."""
    if mode == 'strong':
        return (
            "highpass=f=400:poles=2,"
            "lowpass=f=2800:poles=2,"
            "equalizer=f=1000:t=q:w=1:g=3,"
            "equalizer=f=100:t=q:w=1:g=-10,"
            "equalizer=f=5000:t=q:w=1:g=-10,"
            "compand=attacks=0.1:decays=0.3:points=-80/-80|-45/-30|-20/-15|0/-10|20/-5"
        )
    else:
        return (
            "highpass=f=300:poles=1,"
            "lowpass=f=3400:poles=1,"
            "equalizer=f=1200:t=q:w=2:g=2,"
            "equalizer=f=80:t=q:w=1:g=-6,"
            "equalizer=f=6000:t=q:w=1:g=-6"
        )


@app.route('/api/remove-music', methods=['POST'])
def remove_music():
    """Remove background music from audio or video, keeping human voice."""
    start_time = time.time()
    temp_input = None
    temp_audio_out = None
    temp_video_out = None
    temp_extracted_audio = None
    
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'لم يتم رفع أي ملف'}), 400
        
        media_file = request.files['audio']
        if media_file.filename == '':
            return jsonify({'error': 'اسم الملف فارغ'}), 400
        
        if not allowed_media_clean(media_file.filename):
            return jsonify({'error': 'صيغة الملف غير مدعومة. الصيغ المدعومة: MP3, WAV, M4A, OGG, FLAC, AAC, MP4, MKV, MOV, AVI, WEBM'}), 400
        
        media_file.seek(0, 2)
        file_size = media_file.tell()
        media_file.seek(0)
        
        if file_size > MAX_MEDIA_SIZE_MB_CLEAN * 1024 * 1024:
            return jsonify({'error': f'حجم الملف أكبر من الحد الأقصى المسموح ({MAX_MEDIA_SIZE_MB_CLEAN}MB)'}), 400
        
        mode = request.form.get('mode', 'normal')
        output_type = request.form.get('output_type', 'auto')
        is_video = is_video_file(media_file.filename)
        
        ext = media_file.filename.rsplit('.', 1)[1].lower()
        temp_input = os.path.join(tempfile.gettempdir(), f"media_clean_{uuid.uuid4().hex}.{ext}")
        media_file.save(temp_input)
        
        try:
            probe_cmd = ['ffprobe', '-v', 'error', '-show_entries', 'format=duration', '-of', 'csv=p=0', temp_input]
            probe_result = subprocess.run(probe_cmd, capture_output=True, text=True, timeout=30)
            if probe_result.returncode == 0 and probe_result.stdout.strip():
                duration_seconds = float(probe_result.stdout.strip())
                duration_minutes = duration_seconds / 60
                if duration_minutes > MAX_MEDIA_DURATION_CLEAN:
                    safe_remove_file(temp_input)
                    return jsonify({'error': f'مدة الملف أطول من الحد الأقصى المسموح ({MAX_MEDIA_DURATION_CLEAN} دقيقة)'}), 400
        except Exception as e:
            logging.warning(f"Could not check media duration: {e}")
        
        filter_complex = get_music_removal_filter(mode)
        temp_audio_out = os.path.join(tempfile.gettempdir(), f"audio_cleaned_{uuid.uuid4().hex}.mp3")
        
        if is_video:
            temp_extracted_audio = os.path.join(tempfile.gettempdir(), f"extracted_audio_{uuid.uuid4().hex}.wav")
            extract_cmd = ['ffmpeg', '-i', temp_input, '-vn', '-acodec', 'pcm_s16le', '-ar', '44100', '-ac', '2', '-y', temp_extracted_audio]
            extract_result = subprocess.run(extract_cmd, capture_output=True, text=True, timeout=300)
            
            if extract_result.returncode != 0:
                logging.error(f"FFmpeg extract error: {extract_result.stderr}")
                safe_remove_file(temp_input)
                return jsonify({'error': 'فشل في استخراج الصوت من الفيديو'}), 500
            
            process_cmd = ['ffmpeg', '-i', temp_extracted_audio, '-af', filter_complex, '-acodec', 'libmp3lame', '-ab', '192k', '-ar', '44100', '-y', temp_audio_out]
            process_result = subprocess.run(process_cmd, capture_output=True, text=True, timeout=300)
            
            if process_result.returncode != 0:
                logging.error(f"FFmpeg process error: {process_result.stderr}")
                safe_remove_files(temp_input, temp_extracted_audio)
                return jsonify({'error': 'فشل في معالجة الصوت'}), 500
            
            safe_remove_file(temp_extracted_audio)
            
            if output_type == 'audio_only':
                safe_remove_file(temp_input)
                
                if not os.path.exists(temp_audio_out):
                    return jsonify({'error': 'فشل في إنشاء ملف الصوت المعالج'}), 500
                
                duration_ms = int((time.time() - start_time) * 1000)
                log_activity('remove_music', 'process_video_audio', 'success', duration_ms=duration_ms, file_size=file_size)
                
                original_name = os.path.splitext(media_file.filename)[0]
                
                @after_this_request
                def cleanup_audio(response):
                    safe_remove_file(temp_audio_out)
                    return response
                
                return send_file(temp_audio_out, mimetype='audio/mpeg', as_attachment=True, download_name=f"{original_name}_صوت_نقي.mp3")
            
            else:
                temp_video_out = os.path.join(tempfile.gettempdir(), f"video_cleaned_{uuid.uuid4().hex}.mp4")
                
                merge_cmd = ['ffmpeg', '-i', temp_input, '-i', temp_audio_out, '-c:v', 'copy', '-map', '0:v:0', '-map', '1:a:0', '-shortest', '-y', temp_video_out]
                merge_result = subprocess.run(merge_cmd, capture_output=True, text=True, timeout=600)
                
                if merge_result.returncode != 0:
                    logging.error(f"FFmpeg merge error: {merge_result.stderr}")
                    safe_remove_files(temp_input, temp_audio_out)
                    return jsonify({'error': 'فشل في دمج الفيديو مع الصوت المعالج'}), 500
                
                safe_remove_files(temp_input, temp_audio_out)
                
                if not os.path.exists(temp_video_out):
                    return jsonify({'error': 'فشل في إنشاء ملف الفيديو المعالج'}), 500
                
                duration_ms = int((time.time() - start_time) * 1000)
                log_activity('remove_music', 'process_video', 'success', duration_ms=duration_ms, file_size=file_size)
                
                original_name = os.path.splitext(media_file.filename)[0]
                
                @after_this_request
                def cleanup_video(response):
                    safe_remove_file(temp_video_out)
                    return response
                
                return send_file(temp_video_out, mimetype='video/mp4', as_attachment=True, download_name=f"{original_name}_صوت_نقي.mp4")
        
        else:
            cmd = ['ffmpeg', '-i', temp_input, '-af', filter_complex, '-acodec', 'libmp3lame', '-ab', '192k', '-ar', '44100', '-y', temp_audio_out]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode != 0:
                logging.error(f"FFmpeg error: {result.stderr}")
                safe_remove_file(temp_input)
                return jsonify({'error': 'فشل في معالجة الصوت'}), 500
            
            safe_remove_file(temp_input)
            
            if not os.path.exists(temp_audio_out):
                return jsonify({'error': 'فشل في إنشاء ملف الصوت المعالج'}), 500
            
            duration_ms = int((time.time() - start_time) * 1000)
            log_activity('remove_music', 'process', 'success', duration_ms=duration_ms, file_size=file_size)
            
            original_name = os.path.splitext(media_file.filename)[0]
            
            @after_this_request
            def cleanup(response):
                safe_remove_file(temp_audio_out)
                return response
            
            return send_file(temp_audio_out, mimetype='audio/mpeg', as_attachment=True, download_name=f"{original_name}_صوت_نقي.mp3")
        
    except subprocess.TimeoutExpired:
        safe_remove_files(temp_input, temp_audio_out, temp_video_out, temp_extracted_audio)
        log_activity('remove_music', 'process', 'error', error_message='Timeout')
        return jsonify({'error': 'انتهى الوقت المسموح للمعالجة. جرب ملف أصغر.'}), 408
        
    except Exception as e:
        safe_remove_files(temp_input, temp_audio_out, temp_video_out, temp_extracted_audio)
        log_activity('remove_music', 'process', 'error', error_message=str(e))
        log_error('Exception', str(e), traceback.format_exc(), tool_name='remove_music')
        logging.error(f"Remove music error: {e}")
        return jsonify({'error': f'حدث خطأ: {str(e)}'}), 500


@app.route('/api/video-info', methods=['GET'])
def get_video_info():
    """Get video information from URL (YouTube, TikTok, Instagram, Facebook, etc.)"""
    try:
        url = request.args.get('url', '').strip()
        
        if not url:
            return jsonify({'error': 'الرجاء إدخال رابط فيديو'}), 400
        
        logging.info(f"Fetching video info for: {url}")
        
        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'extract_flat': False,
            'socket_timeout': 30,
            'skip_download': True,
            'no_check_certificates': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            
            # Get thumbnail
            thumbnail = info.get('thumbnail') or info.get('thumbnails', [{}])[-1].get('url', '')
            
            # Get duration
            duration = info.get('duration', 0)
            
            # Get title
            title = info.get('title', 'فيديو بدون عنوان')
            
            # Get uploader/channel
            uploader = info.get('uploader', 'مجهول')
            
            return jsonify({
                'success': True,
                'title': title,
                'duration': duration,
                'thumbnail': thumbnail,
                'uploader': uploader,
                'url': url
            })
    
    except Exception as e:
        logging.error(f"Error fetching video info: {e}")
        return jsonify({
            'success': False,
            'error': 'فشل في جلب معلومات الفيديو. تأكد من الرابط وحاول مرة أخرى.'
        }), 400


@app.route('/api/cut-audio', methods=['POST'])
def cut_audio():
    """Cut audio file - no AI required, just ffmpeg"""
    try:
        start_time = request.form.get('start_time', '0')
        end_time = request.form.get('end_time', '0')
        audio_type = request.form.get('type', 'file')
        
        try:
            start_seconds = float(start_time)
            end_seconds = float(end_time) if end_time else None
        except:
            return jsonify({'error': 'أوقات غير صحيحة'}), 400
        
        if start_seconds < 0:
            return jsonify({'error': 'وقت البداية يجب أن يكون أكبر من 0'}), 400
        
        audio_file = None
        if audio_type == 'file':
            audio_file = request.files.get('audio')
            if not audio_file:
                return jsonify({'error': 'الرجاء رفع ملف صوتي'}), 400
        else:
            url = request.form.get('url')
            if not url:
                return jsonify({'error': 'الرجاء إدخال رابط فيديو'}), 400
            try:
                # Use the existing helper function for consistent YouTube handling
                temp_audio = download_audio_from_youtube(url, tempfile.gettempdir())
                if not temp_audio or not os.path.exists(temp_audio):
                    logging.error(f"YouTube audio file not found")
                    return jsonify({'error': 'فشل في تحويل الفيديو. قد يكون الفيديو محميًا أو غير متاح في منطقتك. جرب رابط مختلف.'}), 400
            except Exception as e:
                logging.error(f"YouTube download error: {e}")
                import traceback
                logging.error(f"Full traceback: {traceback.format_exc()}")
                return jsonify({'error': f'فشل في تحميل الفيديو: تحقق من الرابط وحاول مرة أخرى'}), 400
        
        input_path = None
        if audio_file:
            ext = audio_file.filename.rsplit('.', 1)[-1].lower()
            input_path = os.path.join(tempfile.gettempdir(), f"audio_{uuid.uuid4().hex}.{ext}")
            audio_file.save(input_path)
        else:
            input_path = temp_audio
        
        output_path = os.path.join(tempfile.gettempdir(), f"audio_cut_{uuid.uuid4().hex}.mp3")
        
        ffmpeg_cmd = ['ffmpeg', '-i', input_path, '-ss', str(start_seconds)]
        if end_seconds and end_seconds > start_seconds:
            duration = end_seconds - start_seconds
            ffmpeg_cmd.extend(['-t', str(duration)])
        
        ffmpeg_cmd.extend(['-acodec', 'libmp3lame', '-ab', '192k', '-y', output_path])
        
        result = subprocess.run(ffmpeg_cmd, capture_output=True, text=True, timeout=300)
        
        if result.returncode != 0 or not os.path.exists(output_path):
            logging.error(f"FFmpeg error: {result.stderr}")
            return jsonify({'error': 'فشل في معالجة الملف'}), 500
        
        log_activity('audio_cutter', 'cut_audio', 'success')
        
        @after_this_request
        def cleanup(response):
            safe_remove_file(input_path)
            safe_remove_file(output_path)
            if audio_type == 'url':
                safe_remove_file(temp_audio)
            return response
        
        return send_file(output_path, mimetype='audio/mpeg', as_attachment=True, download_name='audio_cut.mp3')
        
    except subprocess.TimeoutExpired:
        log_activity('audio_cutter', 'cut_audio', 'error', error_message='Timeout')
        return jsonify({'error': 'انتهى الوقت المسموح'}), 408
    except Exception as e:
        log_activity('audio_cutter', 'cut_audio', 'error', error_message=str(e))
        logging.error(f"Audio cutter error: {e}")
        return jsonify({'error': f'حدث خطأ: {str(e)}'}), 500


def init_admin_user():
    """Initialize default admin user if none exists."""
    try:
        admin = AdminUser.query.filter_by(is_admin=True).first()
        if not admin:
            admin = AdminUser(username='admin', is_admin=True)
            admin.set_password('admin123')
            db.session.add(admin)
            db.session.commit()
            logging.info("[Admin] Default admin user created (username: admin, password: admin123)")
            logging.info("[Admin] IMPORTANT: Please change the default password immediately!")
    except Exception as e:
        logging.warning(f"[Admin] Could not initialize admin user: {e}")


with app.app_context():
    try:
        from sqlalchemy import inspect, text
        inspector = inspect(db.engine)
        existing_tables = inspector.get_table_names()
        if not existing_tables:
            try:
                db.create_all()
                logging.info("[DB] Database tables created successfully")
            except Exception as create_error:
                if "already exists" in str(create_error).lower() or "unique" in str(create_error).lower():
                    logging.info("[DB] Tables already exist (created by another worker)")
                else:
                    raise create_error
        else:
            logging.info(f"[DB] Database already initialized with {len(existing_tables)} tables")
        init_admin_user()
        logging.info("=" * 60)
        logging.info(f"Admin dashboard for AI providers is available at: {ADMIN_DASHBOARD_PATH} (requires admin login)")
        logging.info("Default admin credentials: username=admin, password=admin123")
        logging.info(f"Groq provider configured: {ai_manager.groq.is_configured}")
        logging.info(f"HuggingFace provider configured: {ai_manager.huggingface.is_configured}")
        logging.info("=" * 60)
    except Exception as e:
        logging.warning(f"[DB] Initialization handled by another worker or already complete: {e}")


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
